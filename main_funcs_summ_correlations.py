import global_vars
import helper_funcs

import pandas as pd
import numpy as np
from openpyxl import Workbook
from openpyxl.styles import Border, Side, Alignment, Font
from openpyxl.utils import get_column_letter
from openpyxl.utils.dataframe import dataframe_to_rows
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

#-----------------------------------------------------------Output dataframe----------------------------------------------------
def summ_corr_generate_output_df(mod_raw_data_df):
	output_df = mod_raw_data_df.copy()

	return output_df

#-----------------------------------------------------------Saving data----------------------------------------------------
def summ_corr_apa_table(mod_raw_data_df, output_df):
	#unique vars in var1 given my user's order on and ONLY adding vars from col 2 that are NOT in var 1
	variables_list = list(output_df[global_vars.summ_corr_varOne].unique()) + list(set(output_df[global_vars.summ_corr_varTwo].unique()) - set(output_df[global_vars.summ_corr_varOne].unique()))

	wb = Workbook()
	ws = wb.active

	ws.append([""] + variables_list[1:])
	for ind,var in enumerate(variables_list[:-1]):
		ws.cell(row=ind+2, column=1).value = var

	start_col = 2
	for row in range(2, len(variables_list)+1):
		row_var = ws.cell(row=row, column=1).value
		for col in range(start_col, len(variables_list)+1):
			col_var = ws.cell(row=1, column=col).value
			#here query method is not only slower as it is much smaller dataset but also cannot refer to two different vaiables (colname and val)
			df_filtered = output_df[((output_df[global_vars.summ_corr_varOne]==row_var) & (output_df[global_vars.summ_corr_varTwo]==col_var)) | ((output_df[global_vars.summ_corr_varOne]==col_var) & (output_df[global_vars.summ_corr_varTwo]==row_var))].iloc[0]
			r = df_filtered[global_vars.summ_corr_coeff]
			p = df_filtered["adjusted_pvalues"]
			r = helper_funcs.correlations_format_val(r, p)
			ws.cell(row=row, column=col).value = r
		start_col += 1

	for row in range(1, len(variables_list)+1):
		for cell in ws[row]:
			cell.alignment = global_vars.alignment_center

	for cell in ws[1]:
		cell.font = global_vars.font_header

	for row in range(2, len(variables_list)+1):
		ws.cell(row=row, column=1).font = global_vars.font_header

	for cell in ws[len(variables_list)]:
		cell.border = Border(bottom=global_vars.border_APA)

	for cell in ws[1]:
		cell.border = Border(top=global_vars.border_APA, bottom=global_vars.border_APA)

	table_notes = ["**p < 0.01", "*p < {}".format(global_vars.alpha_threshold)]
	helper_funcs.add_table_notes(ws, table_notes)

	helper_funcs.savefile(wb=wb)

def summ_corr_apa_table_word(mod_raw_data_df, output_df):
	#unique vars in var1 given my user's order on and ONLY adding vars from col 2 that are NOT in var 1
	variables_list = list(output_df[global_vars.summ_corr_varOne].unique()) + list(set(output_df[global_vars.summ_corr_varTwo].unique()) - set(output_df[global_vars.summ_corr_varOne].unique()))

	doc = Document()
	table_rows_len, table_cols_len = len(variables_list), len(variables_list)
	table = doc.add_table(rows=table_rows_len, cols=table_cols_len)

	for ind, var in enumerate([""] + variables_list[1:]):
		table.cell(row_idx=0, col_idx=ind).text = var
	for ind, var in enumerate([""] + variables_list[:-1]):
		table.cell(row_idx=ind, col_idx=0).text = var

	start_col = 1
	for row in range(1, table_rows_len):
		row_var = table.cell(row_idx=row, col_idx=0).text
		for col in range(start_col, table_cols_len):
			col_var = table.cell(row_idx=0, col_idx=col).text
			df_filtered = output_df[((output_df[global_vars.summ_corr_varOne]==row_var) & (output_df[global_vars.summ_corr_varTwo]==col_var)) | ((output_df[global_vars.summ_corr_varOne]==col_var) & (output_df[global_vars.summ_corr_varTwo]==row_var))].iloc[0]
			r = df_filtered[global_vars.summ_corr_coeff]
			p = df_filtered["adjusted_pvalues"]
			r = helper_funcs.correlations_format_val(r, p)
			table.cell(row_idx=row, col_idx=col).text = r
		start_col += 1

	for row in range(0, table_rows_len):
		for cell in table.rows[row].cells:
			cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
			cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

	for cell in table.rows[0].cells:
		helper_funcs.word_style(cell, italic=True)
	for cell in table.columns[0].cells:
		helper_funcs.word_style(cell, italic=True)

	for cell in table.rows[0].cells:
		helper_funcs.set_cell_border(cell, top=global_vars.border_APA_word, bottom=global_vars.border_APA_word)
	for cell in table.rows[table_rows_len-1].cells:
		helper_funcs.set_cell_border(cell, bottom=global_vars.border_APA_word)

	doc = helper_funcs.set_autofit(doc)

	p1 = doc.add_paragraph("** ")
	p1.add_run("p").italic = True
	p1.add_run(" < 0.01")
	p2 = doc.add_paragraph("* ")
	p2.add_run("p").italic = True
	p2.add_run(" < {}".format(global_vars.alpha_threshold))
	helper_funcs.add_correction_message_word(doc)

	helper_funcs.savefile(doc=doc)