import pandas as pd
import numpy as np
import global_vars
import helper_funcs

from openpyxl import Workbook
from openpyxl.styles import Border, Side, Alignment, Font
from openpyxl.utils import get_column_letter
from openpyxl.utils.dataframe import dataframe_to_rows
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
#from docx.oxml.table import CT_Row, CT_Tc
#from docx.oxml import OxmlElement
#from docx.oxml.ns import qn
#from docx.table import _Cell
#from docx.shared import Cm
#from docx.shared import Pt


#-----------------------------------------------------------Output dataframe----------------------------------------------------
def raw_corr_generate_output_df(mod_raw_data_df):
	'''
	The critical moment here is that I only have np.nan values where values are non-numeric - can't have anything different from np.nan or a numeric value - bad things happen.
	It's also important to use this type of loop otherwise there is inconsistent behaviour when dropping nan values:
		#https://github.com/scipy/scipy/issues/9103
	Might be worth investigating creating a custom dropping algorithm here as the in-built ones seem to be slow
		#https://github.com/scipy/scipy/issues/6654
	'''
	data_list = []

	i=0
	for var1 in mod_raw_data_df.columns: #can use .columns as all columns are now only of relevant vars (done in previous step)
		for var2 in mod_raw_data_df.columns[i:]:
			if not var1 == var2:
				r, p = helper_funcs.raw_input_corr_coeff(mod_raw_data_df[var1], mod_raw_data_df[var2])
				ci_low, ci_high = helper_funcs.corr_coeff_ci(r, n=min(mod_raw_data_df[var1].count(), mod_raw_data_df[var2].count())) #takes the smallest number as this would be the number of correlations done
				
				data_list.append((var1, var2, r, ci_low, ci_high, p))
		i += 1
	output_df = pd.DataFrame(data_list, columns=["Variable1", "Variable2", "Correlation_Coefficient", "CI_low", "CI_high", "pvalues"])

	return output_df

#-----------------------------------------------------------Saving data----------------------------------------------------
def raw_corr_apa_table(mod_raw_data_df, output_df):
	wb = Workbook()
	ws = wb.active
	
	variables_list = list(mod_raw_data_df.columns) #local var and .columns preferred to minimize use of global scope
	
	ws.append([""] + variables_list[:-1] + ["Mean", "SD"])
	ws["A2"] = variables_list[0]
	for ind,var in enumerate(variables_list[1:]):
		ws.cell(row=(ind*2)+3, column=1).value = var
	
	start_row = 3
	for col in range(2, len(variables_list)+3):
		col_var = ws.cell(row=1, column=col).value
		for row in range(start_row, len(variables_list)*2, 2):
			row_var = ws.cell(row=row, column=1).value
			#Using the query method is slower than using conditionals - in my tests query was between 65 and 80 % slower for 1,000 and 10,000 runs
			r = output_df[((output_df["Variable1"]==row_var) & (output_df["Variable2"]==col_var)) | ((output_df["Variable1"]==col_var) & (output_df["Variable2"]==row_var))].iloc[0]["Correlation_Coefficient"]
			p = output_df[((output_df["Variable1"]==row_var) & (output_df["Variable2"]==col_var)) | ((output_df["Variable1"]==col_var) & (output_df["Variable2"]==row_var))].iloc[0]["adjusted_pvalues"]
			ci_low = output_df[((output_df["Variable1"]==row_var) & (output_df["Variable2"]==col_var)) | ((output_df["Variable1"]==col_var) & (output_df["Variable2"]==row_var))].iloc[0]["CI_low"]
			ci_high = output_df[((output_df["Variable1"]==row_var) & (output_df["Variable2"]==col_var)) | ((output_df["Variable1"]==col_var) & (output_df["Variable2"]==row_var))].iloc[0]["CI_high"]
			r = helper_funcs.correlations_format_val(r, p)
			ci_low, ci_high = helper_funcs.correlations_format_val(ci_low), helper_funcs.correlations_format_val(ci_high)
			ws.cell(row=row, column=col).value = r
			ws.cell(row=row+1, column=col).value = "[" + ci_low + ", " + ci_high + "]"
		start_row += 2

	ws.cell(row=2, column=len(variables_list)+1).value = "{:.2f}".format(mod_raw_data_df[variables_list[0]].mean(skipna=True)) #not a good way of accessing vars; to be redone; see word version below
	ws.cell(row=2, column=len(variables_list)+2).value = "{:.2f}".format(mod_raw_data_df[variables_list[0]].std(skipna=True))
	for row in range(3, len(variables_list)*2, 2):
		ws.cell(row=row, column=len(variables_list)+1).value = "{:.2f}".format(mod_raw_data_df[variables_list[int((row-1)/2)]].mean(skipna=True))
		ws.cell(row=row, column=len(variables_list)+2).value = "{:.2f}".format(mod_raw_data_df[variables_list[int((row-1)/2)]].std(skipna=True))    
	
	for row in range(3, len(variables_list)*2, 2):
		ws.merge_cells(start_row=row, start_column=1, end_row=row+1, end_column=1)
		ws.merge_cells(start_row=row, start_column=len(variables_list)+1, end_row=row+1, end_column=len(variables_list)+1)
		ws.merge_cells(start_row=row, start_column=len(variables_list)+2, end_row=row+1, end_column=len(variables_list)+2)
		
	for row in ws[1:len(variables_list)*2]:
		for cell in row:
			if cell.row > 2 and (cell.column == 1 or cell.column == len(variables_list)+1 or cell.column == len(variables_list)+2):
				cell.alignment = global_vars.alignment_top
			else:
				cell.alignment = global_vars.alignment_center
				
	for cell in ws[1]:
		cell.font = global_vars.font_header
	
	ci_font = Font(size=9)
	for row in range(4, len(variables_list)*2+1, 2):
		for cell in ws[row]:
			cell.font = ci_font
	
	for cell in ws[1]:
		cell.border = Border(top=global_vars.border_APA, bottom=global_vars.border_APA)
	for cell in ws[len(variables_list)*2]:
		cell.border = Border(bottom = global_vars.border_APA)
	
	table_notes = ["Note. M and SD represent mean and standard deviation, respectively. Values in square brackets indicate 95% confidence interval for the correlation coefficient."]
	table_notes.append("Correlation coefficient used: {}".format(list(global_vars.master_dict.keys())[list(global_vars.master_dict.values()).index(global_vars.raw_corr_type)]))
	table_notes.append("**p < 0.01")
	table_notes.append("*p < {}".format(global_vars.alpha_threshold))
	helper_funcs.add_table_notes(ws, table_notes) 
	
	helper_funcs.savefile(wb=wb)

def raw_corr_apa_table_word(mod_raw_data_df, output_df):
	variables_list = list(mod_raw_data_df.columns)

	doc = Document()
	table_rows_len = len(variables_list)*2
	table_cols_len = 1+(len(variables_list)-1)+2 #1 blank cell, all vars but one and mean/sd cols
	table = doc.add_table(rows=table_rows_len, cols=table_cols_len)

	for ind, var in enumerate([""] + variables_list[:-1] + ["Mean", "SD"]):
		table.cell(row_idx=0, col_idx=ind).text = var
	table.cell(row_idx=1, col_idx=0).text = variables_list[0]
	for ind, var in enumerate(variables_list[1:]):
		table.cell(row_idx=(ind*2)+2, col_idx=0).text = var

	start_row = 2
	for col in range(1, table_cols_len):
		col_var = table.cell(row_idx=0, col_idx=col).text
		for row in range(start_row, table_rows_len, 2):
			row_var = table.cell(row_idx=row, col_idx=0).text
			r = output_df[((output_df["Variable1"]==row_var) & (output_df["Variable2"]==col_var)) | ((output_df["Variable1"]==col_var) & (output_df["Variable2"]==row_var))].iloc[0]["Correlation_Coefficient"]
			p = output_df[((output_df["Variable1"]==row_var) & (output_df["Variable2"]==col_var)) | ((output_df["Variable1"]==col_var) & (output_df["Variable2"]==row_var))].iloc[0]["pvalues"]
			ci_low = output_df[((output_df["Variable1"]==row_var) & (output_df["Variable2"]==col_var)) | ((output_df["Variable1"]==col_var) & (output_df["Variable2"]==row_var))].iloc[0]["CI_low"]
			ci_high = output_df[((output_df["Variable1"]==row_var) & (output_df["Variable2"]==col_var)) | ((output_df["Variable1"]==col_var) & (output_df["Variable2"]==row_var))].iloc[0]["CI_high"]
			r = helper_funcs.correlations_format_val(r, p)
			ci_low, ci_high = helper_funcs.correlations_format_val(ci_low), helper_funcs.correlations_format_val(ci_high)
			table.cell(row_idx=row, col_idx=col).text = r
			table.cell(row_idx=row+1, col_idx=col).text = "[" + ci_low + ", " + ci_high + "]"
		start_row += 2

	table.cell(row_idx=1, col_idx=table_cols_len-2).text = "{:.2f}".format(mod_raw_data_df[variables_list[0]].mean(skipna=True))
	table.cell(row_idx=1, col_idx=table_cols_len-1).text = "{:.2f}".format(mod_raw_data_df[variables_list[0]].std(skipna=True))
	for row in range(2, table_rows_len, 2):
		row_var = table.cell(row_idx=row, col_idx=0).text
		table.cell(row_idx=row, col_idx=table_cols_len-2).text = "{:.2f}".format(mod_raw_data_df[row_var].mean(skipna=True))
		table.cell(row_idx=row, col_idx=table_cols_len-1).text = "{:.2f}".format(mod_raw_data_df[row_var].std(skipna=True))

	for row in range(2, table_rows_len, 2):
		table.cell(row_idx=row, col_idx=0).merge(table.cell(row_idx=row+1, col_idx=0))
		table.cell(row_idx=row, col_idx=table_cols_len-2).merge(table.cell(row_idx=row+1, col_idx=table_cols_len-2))
		table.cell(row_idx=row, col_idx=table_cols_len-1).merge(table.cell(row_idx=row+1, col_idx=table_cols_len-1))

	for row in range(0, table_rows_len):
		for col in range(0, table_cols_len):
			table.cell(row_idx=row, col_idx=col).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
			if row > 1 and (col == 0 or col == table_cols_len-2 or col == table_cols_len-1):
				table.cell(row_idx=row, col_idx=col).vertical_alignment = WD_ALIGN_VERTICAL.TOP
			else:
				table.cell(row_idx=row, col_idx=col).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

	for cell in table.rows[0].cells:
		helper_funcs.word_style(cell, italic=True)
	for row in range(3, table_rows_len, 2):
		for col in range(1, table_cols_len):
			helper_funcs.word_style(table.cell(row_idx=row, col_idx=col), size=9)

	for cell in table.rows[0].cells:
		helper_funcs.set_cell_border(cell, top=global_vars.border_APA_word, bottom=global_vars.border_APA_word)
	#below is a workaround for last row as the set_cell_border function does not apply border correctly to merged cells: add temp row and add top border
	temp_row = table.add_row().cells
	for cell in temp_row:
		helper_funcs.set_cell_border(cell, top=global_vars.border_APA_word)

	doc = helper_funcs.set_autofit(doc)

	note = doc.add_paragraph()
	note.add_run("Note").bold = True
	note.add_run(". M and SD represent mean and standard deviation, respectively. Values in square brackets indicate 95% confidence interval for the correlation coefficient.")
	doc.add_paragraph("Correlation coefficient used: {}".format(list(global_vars.master_dict.keys())[list(global_vars.master_dict.values()).index(global_vars.raw_corr_type)]))
	p1 = doc.add_paragraph("** ")
	p1.add_run("p").italic = True
	p1.add_run(" < 0.01")
	p2 = doc.add_paragraph("* ")
	p2.add_run("p").italic = True
	p2.add_run(" < {}".format(global_vars.alpha_threshold))
	helper_funcs.add_correction_message_word(doc)

	helper_funcs.savefile(doc=doc)