# MUltiple tests corrections and FOrmatted tables Software (MUFOS)
# Copyright (C) 2020  Nikolay Petrov, Vasil Atanasov, & Trevor Thompson
import global_vars
import pandas as pd
import numpy as np
from scipy import stats
from openpyxl import Workbook
from openpyxl.utils.dataframe import dataframe_to_rows
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import _Cell
from docx.shared import Pt

#----------------------------------------------Helper functions for GUI
def get_df_columns(path_and_filename):
	try:
		if path_and_filename.endswith(".xlsx"):
			df_columns = list(pd.read_excel(path_and_filename, sheet_name=0).columns)
		elif path_and_filename.endswith(".csv"):
			df_columns = list(pd.read_csv(path_and_filename).columns)

		return df_columns
	except:
		raise Exception("There was a problem reading the columns from the file.")

def get_raw_indttest_grouplevels(path_and_filename, group_var):
	try:
		if path_and_filename.endswith(".xlsx"):
			df = pd.read_excel(path_and_filename, sheet_name=0)
		elif path_and_filename.endswith(".csv"):
			df = pd.read_csv(path_and_filename)
	except:
		raise Exception("There was a problem reading the columns from the file.")

	unique_groups = [x.strip() for x in df[group_var].astype(str).unique() if len(x.strip()) > 0]

	return unique_groups

#-----------------------------------------------Helper functions fo generating modified raw data dataframes
def raw_input_generate_mod_raw_data_df(raw_data_df, numeric_cols, non_numeric_cols=[]):
	# reminder: don't drop na; just coerce
	# function is also used for summ_corr input type
	mod_raw_data_df = raw_data_df[numeric_cols + non_numeric_cols] # does NOT raise errors when array is blank (default arg)
	mod_raw_data_df[numeric_cols] = mod_raw_data_df[numeric_cols].apply(pd.to_numeric, errors="coerce") # non-numeric values will be np.nan
	try:
		mod_raw_data_df[non_numeric_cols] = mod_raw_data_df[non_numeric_cols].astype(str) # does NOT raise errors when array is blank (default arg)
	except:
		raise Exception("The data could not be processed. Please ensure that the non-numeric columns contain only strings.")
	
	return mod_raw_data_df

def error_on_input(df, cols, input_type):
	'''
	df - the dataframe to check
	cols - the cols to check
	input_type - the input type to check for - can be either 'numeric' or 'nan'
	'''
	#creating a dictionary - key is col name value is a list of all nan values
	bad_vals_dict = {}
	for var in cols:
		if input_type == "numeric":
			current_series = pd.to_numeric(df[var], errors="coerce")
		elif input_type == "nan":
			current_series = df[var]
		ind_list = list(current_series.isnull().to_numpy().nonzero()[0])
		bad_vals_dict[var] = ind_list
	
	# in pairttest, it is possible to have 1 pair of vars of length X and another of X+200. To handle that, the final non-nan case is found for each pair
	# and a dictionary is created where the key is the col name and the value is the index of the final case
	if global_vars.raw_test == "pairttest":
		pairttest_final_cases_ind_dict = {}
		for ind in range(0, len(cols), 2):
			pair_final_case = max(df[df.columns[ind]].dropna().index[-1], df[df.columns[ind+1]].dropna().index[-1]) #max here of the two vars as final case can have missing data
			pairttest_final_cases_ind_dict[df.columns[ind]] = pair_final_case
			pairttest_final_cases_ind_dict[df.columns[ind+1]] = pair_final_case
		bad_vals_dict = {k: [e for e in v if e <= pairttest_final_cases_ind_dict[k]] for k, v in bad_vals_dict.items()}
	
	bad_vals_dict = {k: v for k, v in bad_vals_dict.items() if len(v) > 0} # dictionary reset to contain only lists of missing data as values, if none function ends
	if bad_vals_dict != {}:
		error_msg = "Non-numerical or blank entries found in the data!\n\n"
		for col, ind_arr in bad_vals_dict.items():
			ind_str = (", ").join([str(val+2) for counter, val in enumerate(ind_arr) if counter < 10])
			if len(ind_arr) > 10:
				ind_str += "... and {} more".format(len(ind_arr)-10)
			error_msg += "In column \"{c}\", there are non-numerical/blank entries on the following rows: {i}\n".format(c=col, i=ind_str)
		raise Exception(error_msg)

#------------------------------------------------Helper functions for generating output dataframes
def raw_input_corr_coeff(x, y):
	if global_vars.raw_corr_type == "pearson":
		# although it's a clunky loop, there does not seem to be another way around:
		# https://stackoverflow.com/questions/38894488/dropping-nan-with-pearsons-r-in-scipy-pandas
		x = np.array(x)
		y = np.array(y)
		nas = ~np.logical_or(np.isnan(x), np.isnan(y))
		x = np.compress(nas, x) # can't just ignore NAN; compress better here as it allows both arrays to have same length which is a must for stats.pearsonr()
		y = np.compress(nas, y)

		r, p = stats.pearsonr(x, y)
	elif global_vars.raw_corr_type == "spearman":
		r, p = stats.spearmanr(x, y, nan_policy="omit")
	elif global_vars.raw_corr_type == "kendall":
		r, p = stats.kendalltau(x, y, nan_policy="omit")

	return r, p

def corr_coeff_ci(r, n):
	'''
	Credits: https://zhiyzuo.github.io/Pearson-Correlation-CI-in-Python/
	'''
	if n < 4:
		raise Exception("Sample size is too low for correlations. Minimum of 4 observations per variable required.")

	r_z = np.arctanh(r)
	se = 1/np.sqrt(n-3)
	z = stats.norm.ppf(1-0.05/2)
	low_z = r_z - z*se 
	high_z = r_z + z*se
	low, high = np.tanh((low_z, high_z))
	
	return low, high

def calc_ttest_effect_size(effect_size_choice, t=None, n1=None, n2=None, m1=None, m2=None, sd1=None):
	if effect_size_choice == "Cohen's d":
		effect_size = t * np.sqrt(1 / n1 + 1 / n2)
	elif effect_size_choice == "Hedge's g":
		# calculated using cohens_d * (1 - (3/4(n1+n2-9))) where cohens d is t*sqrt(1/n1 + 1/n2)
		effect_size = (t * np.sqrt(1 / n1 + 1 / n2)) * (1 - (3 / (4 * (n1 + n2 - 9))))
	elif effect_size_choice == "Glass's delta":
		# glass delta = (Mean2 - Mean1) / SD1 where SD1 is always that of control
		effect_size = (m2 - m1) / sd1
	elif effect_size_choice == "None":
		effect_size = np.nan

	return effect_size

#----------------------------------------------------Helper functions for apa tables
def pvalue_formatting(x):
	if x<0.001:
		return "<.001"
	else:
		return "{:.3f}".format(x).replace("0", "", 1)

def correlations_format_val(x, p=None):
	if isinstance(x, str): # covers the case when the correlations come from SPSS input and are flagged significant
		try:
			x = float("".join([char for char in x if char!="*"]))
		except:
			raise Exception("Something went wrong with formatting the correlation coefficients. Please try again.")
		
	if abs(x) == 1.0:
		x = "{:.2f}".format(val)
	else:
		x = "{:.2f}".format(x).replace("0","",1)
	
	if p != None:
		#adding "_" for Word so I can split in the function and add the **s in a different col for formatting purposes
		if p < 0.001:
			x = x + "_**" if global_vars.output_filetype == "Word" else x + "**"
		elif p < 0.05:
			x = x + "_*" if global_vars.output_filetype == "Word" else x + "*"
	return x

def savefile(doc=None, wb=None):
	try:
		if doc != None:
			doc.save(global_vars.output_filename + ".docx")
		elif wb != None:
			wb.save(filename=global_vars.output_filename + ".xlsx")
	except PermissionError:
		raise Exception("The file could not be saved. Please check if a file with the same name already exists in the target directory and close/delete it; then try running the software again.")

#-------------Excel
def add_table_notes(ws, table_notes=[]):
	if global_vars.correction_type != "None":
		#clunky expression but it's a one-off
		table_notes.append("Multiple tests correction applied to p values: {c}".format(c=list(global_vars.master_dict.keys())[list(global_vars.master_dict.values()).index(global_vars.correction_type)]))
	for note in table_notes:
		ws.append([note])

#-------------Word
def add_correction_message_word(doc):
	if global_vars.correction_type != "None":
		#clunky expression but it's a one-off
		doc.add_paragraph("Multiple tests correction applied to p values: {c}".format(c=list(global_vars.master_dict.keys())[list(global_vars.master_dict.values()).index(global_vars.correction_type)]))

def word_style(cell, italic=False, bold=False, size=None):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            if italic:
                run.font.italic = True
            if bold:
                run.font.bold = True
            if size != None:
                run.font.size = Pt(size)
                
def set_autofit(doc: Document) -> Document:
    #credits - https://github.com/python-openxml/python-docx/issues/209#issuecomment-566128709
    """
    Hotfix for autofit.
    """
    for t_idx, table in enumerate(doc.tables):
        doc.tables[t_idx].autofit = True
        doc.tables[t_idx].allow_autofit = True
        doc.tables[t_idx]._tblPr.xpath("./w:tblW")[0].attrib["{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type"] = "auto"
        for row_idx, r_val in enumerate(doc.tables[t_idx].rows):
            for cell_idx, c_val in enumerate(doc.tables[t_idx].rows[row_idx].cells):
                doc.tables[t_idx].rows[row_idx].cells[cell_idx]._tc.tcPr.tcW.type = 'auto'
                doc.tables[t_idx].rows[row_idx].cells[cell_idx]._tc.tcPr.tcW.w = 0
    return doc

def set_cell_border(cell: _Cell, **kwargs):
    #credits - https://stackoverflow.com/a/49615968/13078832
    """
    Set cell`s border
    Usage:

    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def delete_columns_word(table, columns):
	# does not work around merged cells
	# credits - https://github.com/python-openxml/python-docx/issues/441#issuecomment-498716546
	# sort columns descending
	columns.sort(reverse=True)
	
	grid = table._tbl.find("w:tblGrid", table._tbl.nsmap)
	for ci in columns:
		for cell in table.column_cells(ci):
			cell._tc.getparent().remove(cell._tc)

		# Delete column reference.
		col_elem = grid[ci]
		grid.remove(col_elem)