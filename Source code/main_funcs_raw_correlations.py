# MUltiple tests corrections and FOrmatted tables Software (MUFOS)
# Copyright (C) 2020  Nikolay Petrov, Vasil Atanasov, & Trevor Thompson
import pandas as pd
import numpy as np
import global_vars
import helper_funcs

from openpyxl import Workbook
from openpyxl.styles import Border, Side, Alignment, Font
from openpyxl.utils import get_column_letter
from openpyxl.utils.dataframe import dataframe_to_rows
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Inches


#-----------------------------------------------------------Output dataframe----------------------------------------------------
def raw_corr_generate_output_df(mod_raw_data_df):
	'''
	The critical moment here is that I only have np.nan values where values are non-numeric - can't have anything different from np.nan or a numeric value - bad things happen.
	It's also important to use this type of loop otherwise there is inconsistent behaviour when dropping nan values: https://github.com/scipy/scipy/issues/9103
	Might be worth investigating creating a custom dropping algorithm here as the in-built ones seem to be slow: https://github.com/scipy/scipy/issues/6654
	'''
	data_list = []

	i=0
	for var1 in mod_raw_data_df.columns: # can use .columns as all columns are now only of relevant vars (done in previous step)
		for var2 in mod_raw_data_df.columns[i:]:
			if not var1 == var2:
				r, p = helper_funcs.raw_input_corr_coeff(mod_raw_data_df[var1], mod_raw_data_df[var2])
				ci_low, ci_high = helper_funcs.corr_coeff_ci(r, n=min(mod_raw_data_df[var1].count(), mod_raw_data_df[var2].count())) # takes the smallest number as this would be the number of correlations done
				
				data_list.append((var1, var2, r, ci_low, ci_high, p))
		i += 1
	output_df = pd.DataFrame(data_list, columns=["Variable1", "Variable2", "Correlation_Coefficient", "CI_low", "CI_high", "pvalues"])

	return output_df

#-----------------------------------------------------------Saving data----------------------------------------------------
def raw_corr_apa_table_noCIs_excel(mod_raw_data_df, output_df):
	variables_list = list(mod_raw_data_df.columns) #local var and .columns preferred to minimize use of global scope

	wb = Workbook()
	ws = wb.active
	
	if global_vars.corr_table_triangle == "Upper triangle":
		header_cols = variables_list[1:]
		header_rows = variables_list[:-1]
	elif global_vars.corr_table_triangle == "Lower triangle":
		header_cols = variables_list[:-1]
		header_rows = variables_list[1:]
	elif global_vars.corr_table_triangle == "Both":
		header_cols = [x for x in variables_list]
		header_rows = [x for x in variables_list]
	
	ws.append([""] + header_cols)
	for ind,var in enumerate(header_rows):
		ws.cell(row=ind+2, column=1).value = var

	inside_loop_ind_start = 2
	for outside_loop_ind in range(2, len(header_rows) + 2):

		if global_vars.corr_table_triangle == "Upper triangle" or global_vars.corr_table_triangle == "Both":
			outside_loop_var = ws.cell(row=outside_loop_ind, column=1).value
		elif global_vars.corr_table_triangle == "Lower triangle":
			outside_loop_var = ws.cell(row=1, column=outside_loop_ind).value
		
		for inside_loop_ind in range(inside_loop_ind_start, len(header_cols) + 2):
			
			if global_vars.corr_table_triangle == "Upper triangle" or global_vars.corr_table_triangle == "Both":
				inside_loop_var = ws.cell(row=1, column=inside_loop_ind).value
			elif global_vars.corr_table_triangle == "Lower triangle":
				inside_loop_var = ws.cell(row=inside_loop_ind, column=1).value

			if outside_loop_var == inside_loop_var:
				ws.cell(row=outside_loop_ind, column=inside_loop_ind).value = 1
			else:
				# here query method is not preferred as it is not only slower as it is much smaller dataset but also cannot refer to two different vaiables (colname and val)
				df_filtered = output_df[((output_df["Variable1"]==outside_loop_var) & (output_df["Variable2"]==inside_loop_var)) | ((output_df["Variable1"]==inside_loop_var) & (output_df["Variable2"]==outside_loop_var))].iloc[0]
				r = df_filtered["Correlation_Coefficient"]
				p = df_filtered["adjusted_pvalues"]
				r = helper_funcs.correlations_format_val(r, p)

				if global_vars.corr_table_triangle == "Upper triangle" or global_vars.corr_table_triangle == "Both":
					ws.cell(row=outside_loop_ind, column=inside_loop_ind).value = r
				elif global_vars.corr_table_triangle == "Lower triangle":
					ws.cell(row=inside_loop_ind, column=outside_loop_ind).value = r
		
		if global_vars.corr_table_triangle != "Both":
			inside_loop_ind_start += 1	
		
	for row in range(1, len(header_rows) + 2):
		for cell in ws[row]:
			cell.alignment = global_vars.alignment_center

	for cell in ws[1] + ws["A"]:
		cell.font = global_vars.font_header

	for cell in ws[1]:
		cell.border = Border(top=global_vars.border_APA, bottom=global_vars.border_APA)
	for cell in ws[len(header_rows) + 1]:
		cell.border = Border(bottom = global_vars.border_APA)
	
	table_notes = ["Correlation coefficient used: {}".format(list(global_vars.master_dict.keys())[list(global_vars.master_dict.values()).index(global_vars.raw_corr_type)])]
	table_notes.append("**p < 0.01")
	table_notes.append("*p < {}".format(global_vars.alpha_threshold))
	helper_funcs.add_table_notes(ws, table_notes) 

	helper_funcs.savefile(wb=wb)

def raw_corr_apa_table_withCIs_excel(mod_raw_data_df, output_df):
	variables_list = list(mod_raw_data_df.columns) # local var and .columns preferred to minimize use of global scope

	wb = Workbook()
	ws = wb.active
	
	if global_vars.corr_table_triangle == "Upper triangle":
		header_cols = variables_list[1:]
		header_rows = variables_list[:-1]
	elif global_vars.corr_table_triangle == "Lower triangle":
		header_cols = variables_list[:-1]
		header_rows = variables_list[1:]
	elif global_vars.corr_table_triangle == "Both":
		header_cols = [x for x in variables_list]
		header_rows = [x for x in variables_list]
	# this adds an empty row where the CIs will be
	# the code looks slightly wrong as it return a None array but it works
	[header_rows.insert(x, "") for x in range(1, len(header_rows)*2, 2)]

	ws.append([""] + header_cols)
	for ind,var in enumerate(header_rows):
		ws.cell(row=ind+2, column=1).value = var

	if global_vars.corr_table_triangle == "Upper triangle":
		inside_loop_ind_start = 2
		for outside_loop_ind in range(2, len(header_rows) + 2, 2):
			outside_loop_var = ws.cell(row=outside_loop_ind, column=1).value
			for inside_loop_ind in range(inside_loop_ind_start, len(header_cols) + 2):
				inside_loop_var = ws.cell(row=1, column=inside_loop_ind).value
				# here query method is not preferred as it is not only slower as it is much smaller dataset but also cannot refer to two different vaiables (colname and val)
				df_filtered = output_df[((output_df["Variable1"]==outside_loop_var) & (output_df["Variable2"]==inside_loop_var)) | ((output_df["Variable1"]==inside_loop_var) & (output_df["Variable2"]==outside_loop_var))].iloc[0]
				r = df_filtered["Correlation_Coefficient"]
				p = df_filtered["adjusted_pvalues"]
				r = helper_funcs.correlations_format_val(r, p)
				ci_low, ci_high = helper_funcs.correlations_format_val(df_filtered["CI_low"]), helper_funcs.correlations_format_val(df_filtered["CI_high"])
				ws.cell(row=outside_loop_ind, column=inside_loop_ind).value = r
				ws.cell(row=outside_loop_ind+1, column=inside_loop_ind).value = "[" + ci_low + ", " + ci_high + "]"
			inside_loop_ind_start += 1
	elif global_vars.corr_table_triangle == "Lower triangle":
		inside_loop_ind_start = 2
		for outside_loop_ind in range(2, len(header_cols) + 2):
			outside_loop_var = ws.cell(row=1, column=outside_loop_ind).value
			for inside_loop_ind in range(inside_loop_ind_start, len(header_rows) + 2, 2):
				inside_loop_var = ws.cell(row=inside_loop_ind, column=1).value
				# here query method is not preferred as it is not only slower as it is much smaller dataset but also cannot refer to two different vaiables (colname and val)
				df_filtered = output_df[((output_df["Variable1"]==outside_loop_var) & (output_df["Variable2"]==inside_loop_var)) | ((output_df["Variable1"]==inside_loop_var) & (output_df["Variable2"]==outside_loop_var))].iloc[0]
				r = df_filtered["Correlation_Coefficient"]
				p = df_filtered["adjusted_pvalues"]
				r = helper_funcs.correlations_format_val(r, p)
				ci_low, ci_high = helper_funcs.correlations_format_val(df_filtered["CI_low"]), helper_funcs.correlations_format_val(df_filtered["CI_high"])
				ws.cell(row=inside_loop_ind, column=outside_loop_ind).value = r
				ws.cell(row=inside_loop_ind+1, column=outside_loop_ind).value = "[" + ci_low + ", " + ci_high + "]"
			inside_loop_ind_start += 2	
	elif global_vars.corr_table_triangle == "Both":
		inside_loop_ind_start = 2
		for outside_loop_ind in range(2, len(header_cols) + 2):
			outside_loop_var = ws.cell(row=1, column=outside_loop_ind).value
			for inside_loop_ind in range(inside_loop_ind_start, len(header_rows) + 2, 2):
				inside_loop_var = ws.cell(row=inside_loop_ind, column=1).value
				if outside_loop_var == inside_loop_var:
					ws.cell(row=inside_loop_ind, column=outside_loop_ind).value = "1"
				else:
					# here query method is not preferred as it is not only slower as it is much smaller dataset but also cannot refer to two different vaiables (colname and val)
					df_filtered = output_df[((output_df["Variable1"]==outside_loop_var) & (output_df["Variable2"]==inside_loop_var)) | ((output_df["Variable1"]==inside_loop_var) & (output_df["Variable2"]==outside_loop_var))].iloc[0]
					r = df_filtered["Correlation_Coefficient"]
					p = df_filtered["adjusted_pvalues"]
					r = helper_funcs.correlations_format_val(r, p)
					ci_low, ci_high = helper_funcs.correlations_format_val(df_filtered["CI_low"]), helper_funcs.correlations_format_val(df_filtered["CI_high"])
					ws.cell(row=inside_loop_ind, column=outside_loop_ind).value = r
					ws.cell(row=inside_loop_ind+1, column=outside_loop_ind).value = "[" + ci_low + ", " + ci_high + "]"

	for row in range(2, len(header_rows)+1, 2):
		ws.merge_cells(start_row=row, start_column=1, end_row=row+1, end_column=1)
	if global_vars.corr_table_triangle == "Both":
		for col in range(2, len(header_cols)+2):
			ws.merge_cells(start_row=col*2-2, start_column=col, end_row=col*2-1, end_column=col)

	for row in range(1, len(header_rows) + 2):
		for cell in ws[row]:
			cell.alignment = global_vars.alignment_center

	for cell in ws[1] + ws["A"]:
		cell.font = global_vars.font_header

	for row in range(3, len(header_rows)+2, 2):
		for cell in ws[row]:
			cell.font = Font(size=9)

	for cell in ws[1]:
		cell.border = Border(top=global_vars.border_APA, bottom=global_vars.border_APA)
	for cell in ws[len(header_rows) + 1]:
		cell.border = Border(bottom = global_vars.border_APA)
	
	table_notes = ["Correlation coefficient used: {}".format(list(global_vars.master_dict.keys())[list(global_vars.master_dict.values()).index(global_vars.raw_corr_type)])]
	table_notes.append("**p < 0.01")
	table_notes.append("*p < {}".format(global_vars.alpha_threshold))
	helper_funcs.add_table_notes(ws, table_notes) 

	helper_funcs.savefile(wb=wb)

def raw_corr_apa_table_noCIs_word(mod_raw_data_df, output_df):
	variables_list = list(mod_raw_data_df.columns)

	if global_vars.corr_table_triangle == "Upper triangle":
		header_cols = variables_list[1:]
		header_rows = variables_list[:-1]
	elif global_vars.corr_table_triangle == "Lower triangle":
		header_cols = variables_list[:-1]
		header_rows = variables_list[1:]
	elif global_vars.corr_table_triangle == "Both":
		header_cols = [x for x in variables_list]
		header_rows = [x for x in variables_list]
	# this adds an empty column where the significance signs will be placed for better presentation
	# the code looks slightly wrong as it return a None array but it works
	[header_cols.insert(x, "") for x in range(1, len(header_cols)*2, 2)]

	doc = Document()
	table_rows_len = len(header_rows) + 1
	table_cols_len = len(header_cols) + 1
	table = doc.add_table(rows=table_rows_len, cols=table_cols_len)

	for ind, var in enumerate([""] + header_cols):
		table.cell(row_idx=0, col_idx=ind).text = var
	for ind, var in enumerate([""] + header_rows):
		table.cell(row_idx=ind, col_idx=0).text = var

	if global_vars.corr_table_triangle == "Upper triangle":
		inside_loop_ind_start = 1
		for outside_loop_ind in range(1, table_rows_len):
			outside_loop_var = table.cell(row_idx=outside_loop_ind, col_idx=0).text
			for inside_loop_ind in range(inside_loop_ind_start, table_cols_len, 2):
				inside_loop_var = table.cell(row_idx=0, col_idx=inside_loop_ind).text
				# here query method is not preferred as it is not only slower as it is much smaller dataset but also cannot refer to two different vaiables (colname and val)
				df_filtered = output_df[((output_df["Variable1"]==outside_loop_var) & (output_df["Variable2"]==inside_loop_var)) | ((output_df["Variable1"]==inside_loop_var) & (output_df["Variable2"]==outside_loop_var))].iloc[0]
				r = df_filtered["Correlation_Coefficient"]
				p = df_filtered["adjusted_pvalues"]
				r = helper_funcs.correlations_format_val(r, p)
				if "_" in r:
					r, sign = r.split("_")
				else:
					sign = ""
				table.cell(row_idx=outside_loop_ind, col_idx=inside_loop_ind).text = r
				table.cell(row_idx=outside_loop_ind, col_idx=inside_loop_ind+1).text = sign
			inside_loop_ind_start += 2
	elif global_vars.corr_table_triangle == "Lower triangle":
		inside_loop_ind_start = 1
		for outside_loop_ind in range(1, table_cols_len, 2):
			outside_loop_var = table.cell(row_idx=0, col_idx=outside_loop_ind).text
			for inside_loop_ind in range(inside_loop_ind_start, table_rows_len):
				inside_loop_var = table.cell(row_idx=inside_loop_ind, col_idx=0).text
				# here query method is not preferred as it is not only slower as it is much smaller dataset but also cannot refer to two different vaiables (colname and val)
				df_filtered = output_df[((output_df["Variable1"]==outside_loop_var) & (output_df["Variable2"]==inside_loop_var)) | ((output_df["Variable1"]==inside_loop_var) & (output_df["Variable2"]==outside_loop_var))].iloc[0]
				r = df_filtered["Correlation_Coefficient"]
				p = df_filtered["adjusted_pvalues"]
				r = helper_funcs.correlations_format_val(r, p)
				if "_" in r:
					r, sign = r.split("_")
				else:
					sign = ""
				table.cell(row_idx=inside_loop_ind, col_idx=outside_loop_ind).text = r
				table.cell(row_idx=inside_loop_ind, col_idx=outside_loop_ind+1).text = sign
			inside_loop_ind_start += 1
	elif global_vars.corr_table_triangle == "Both":
		inside_loop_ind_start = 1
		for outside_loop_ind in range(1, table_cols_len, 2):
			outside_loop_var = table.cell(row_idx=0, col_idx=outside_loop_ind).text
			for inside_loop_ind in range(inside_loop_ind_start, table_rows_len):
				inside_loop_var = table.cell(row_idx=inside_loop_ind, col_idx=0).text
				if outside_loop_var == inside_loop_var:
					table.cell(row_idx=inside_loop_ind, col_idx=outside_loop_ind).text = "1"
				else:
				# here query method is not preferred as it is not only slower as it is much smaller dataset but also cannot refer to two different vaiables (colname and val)
					df_filtered = output_df[((output_df["Variable1"]==outside_loop_var) & (output_df["Variable2"]==inside_loop_var)) | ((output_df["Variable1"]==inside_loop_var) & (output_df["Variable2"]==outside_loop_var))].iloc[0]
					r = df_filtered["Correlation_Coefficient"]
					p = df_filtered["adjusted_pvalues"]
					r = helper_funcs.correlations_format_val(r, p)
					if "_" in r:
						r, sign = r.split("_")
					else:
						sign = ""
					table.cell(row_idx=inside_loop_ind, col_idx=outside_loop_ind).text = r
					table.cell(row_idx=inside_loop_ind, col_idx=outside_loop_ind+1).text = sign


	for cell in table.rows[0].cells:
		helper_funcs.set_cell_border(cell, top=global_vars.border_APA_word, bottom=global_vars.border_APA_word)
	for cell in table.rows[table_rows_len-1].cells:
		helper_funcs.set_cell_border(cell, bottom=global_vars.border_APA_word)

	for i in range(1, table_cols_len, 2):
		table.cell(row_idx=0, col_idx=i).merge(table.cell(row_idx=0, col_idx=i+1))
		table.cell(row_idx=0, col_idx=i).text = table.cell(row_idx=0, col_idx=i).text[:-1] # meging cells adds a paragraph break at the end; this gets rid of it

	for cell in table.rows[0].cells:
		helper_funcs.word_style(cell, italic=True)
		cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
	for cell in table.columns[0].cells:
		helper_funcs.word_style(cell, italic=True)
		cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

	for row in range(0, table_rows_len):
		for cell in table.rows[row].cells:
			cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
			cell.paragraphs[0].paragraph_format.space_after = Inches(0)
			cell.paragraphs[0].paragraph_format.space_before = Inches(0)
			cell.paragraphs[0].paragraph_format.line_spacing = 1

	for row in range(1, table_rows_len):
		for col in range(1, table_cols_len, 2):
			table.cell(row_idx=row, col_idx=col).paragraphs[0].paragraph_format.right_indent = Inches(-0.08)
			table.cell(row_idx=row, col_idx=col+1).paragraphs[0].paragraph_format.left_indent = Inches(-0.06)
			table.cell(row_idx=row, col_idx=col).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
			table.cell(row_idx=row, col_idx=col+1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

	doc = helper_funcs.set_autofit(doc)

	para = doc.add_paragraph("** ")
	para.add_run("p").italic = True
	para.add_run(" < 0.01")
	para.add_run().add_break()
	para.add_run("* ")
	para.add_run("p").italic = True
	para.add_run(" < {}".format(global_vars.alpha_threshold))
	doc.add_paragraph("Correlation coefficient used: {}".format(list(global_vars.master_dict.keys())[list(global_vars.master_dict.values()).index(global_vars.raw_corr_type)]))
	helper_funcs.add_correction_message_word(doc)
	
	helper_funcs.savefile(doc=doc)

def raw_corr_apa_table_withCIs_word(mod_raw_data_df, output_df):
	variables_list = list(mod_raw_data_df.columns)

	if global_vars.corr_table_triangle == "Upper triangle":
		header_cols = variables_list[1:]
		header_rows = variables_list[:-1]
	elif global_vars.corr_table_triangle == "Lower triangle":
		header_cols = variables_list[:-1]
		header_rows = variables_list[1:]
	elif global_vars.corr_table_triangle == "Both":
		header_cols = [x for x in variables_list]
		header_rows = [x for x in variables_list]
	# this adds an empty column where the significance signs will be placed for better presentation
	# then adds empty rows for the CIs
	# the code looks slightly wrong as it return a None array but it works
	[header_cols.insert(x, "") for x in range(1, len(header_cols)*2, 2)]
	[header_rows.insert(x, "") for x in range(1, len(header_rows)*2, 2)]

	doc = Document()
	table_rows_len = len(header_rows) + 1
	table_cols_len = len(header_cols) + 1
	table = doc.add_table(rows=table_rows_len, cols=table_cols_len)

	for ind, var in enumerate([""] + header_cols):
		table.cell(row_idx=0, col_idx=ind).text = var
	for ind, var in enumerate([""] + header_rows):
		table.cell(row_idx=ind, col_idx=0).text = var

	if global_vars.corr_table_triangle == "Upper triangle":
		inside_loop_ind_start = 1
		for outside_loop_ind in range(1, table_rows_len, 2):
			outside_loop_var = table.cell(row_idx=outside_loop_ind, col_idx=0).text
			for inside_loop_ind in range(inside_loop_ind_start, table_cols_len, 2):
				inside_loop_var = table.cell(row_idx=0, col_idx=inside_loop_ind).text
				# here query method is not preferred as it is not only slower as it is much smaller dataset but also cannot refer to two different vaiables (colname and val)
				df_filtered = output_df[((output_df["Variable1"]==outside_loop_var) & (output_df["Variable2"]==inside_loop_var)) | ((output_df["Variable1"]==inside_loop_var) & (output_df["Variable2"]==outside_loop_var))].iloc[0]
				r = df_filtered["Correlation_Coefficient"]
				p = df_filtered["adjusted_pvalues"]
				ci_low, ci_high = helper_funcs.correlations_format_val(df_filtered["CI_low"]), helper_funcs.correlations_format_val(df_filtered["CI_high"])
				r = helper_funcs.correlations_format_val(r, p)
				if "_" in r:
					r, sign = r.split("_")
				else:
					sign = ""
				table.cell(row_idx=outside_loop_ind, col_idx=inside_loop_ind).text = r
				table.cell(row_idx=outside_loop_ind, col_idx=inside_loop_ind+1).text = sign
				table.cell(row_idx=outside_loop_ind+1, col_idx=inside_loop_ind).text = "[" + ci_low + ", " + ci_high + "]"
			inside_loop_ind_start += 2
	elif global_vars.corr_table_triangle == "Lower triangle":
		inside_loop_ind_start = 1
		for outside_loop_ind in range(1, table_cols_len, 2):
			outside_loop_var = table.cell(row_idx=0, col_idx=outside_loop_ind).text
			for inside_loop_ind in range(inside_loop_ind_start, table_rows_len, 2):
				inside_loop_var = table.cell(row_idx=inside_loop_ind, col_idx=0).text
				# here query method is not preferred as it is not only slower as it is much smaller dataset but also cannot refer to two different vaiables (colname and val)
				df_filtered = output_df[((output_df["Variable1"]==outside_loop_var) & (output_df["Variable2"]==inside_loop_var)) | ((output_df["Variable1"]==inside_loop_var) & (output_df["Variable2"]==outside_loop_var))].iloc[0]
				r = df_filtered["Correlation_Coefficient"]
				p = df_filtered["adjusted_pvalues"]
				ci_low, ci_high = helper_funcs.correlations_format_val(df_filtered["CI_low"]), helper_funcs.correlations_format_val(df_filtered["CI_high"])
				r = helper_funcs.correlations_format_val(r, p)
				if "_" in r:
					r, sign = r.split("_")
				else:
					sign = ""
				table.cell(row_idx=inside_loop_ind, col_idx=outside_loop_ind).text = r
				table.cell(row_idx=inside_loop_ind, col_idx=outside_loop_ind+1).text = sign
				table.cell(row_idx=inside_loop_ind+1, col_idx=outside_loop_ind).text = "[" + ci_low + ", " + ci_high + "]"
			inside_loop_ind_start += 2
	elif global_vars.corr_table_triangle == "Both":
		inside_loop_ind_start = 1
		for outside_loop_ind in range(1, table_cols_len, 2):
			outside_loop_var = table.cell(row_idx=0, col_idx=outside_loop_ind).text
			for inside_loop_ind in range(inside_loop_ind_start, table_rows_len, 2):
				inside_loop_var = table.cell(row_idx=inside_loop_ind, col_idx=0).text
				if outside_loop_var == inside_loop_var:
					table.cell(row_idx=inside_loop_ind, col_idx=outside_loop_ind).text = "1"
				else:
				# here query method is not preferred as it is not only slower as it is much smaller dataset but also cannot refer to two different vaiables (colname and val)
					df_filtered = output_df[((output_df["Variable1"]==outside_loop_var) & (output_df["Variable2"]==inside_loop_var)) | ((output_df["Variable1"]==inside_loop_var) & (output_df["Variable2"]==outside_loop_var))].iloc[0]
					r = df_filtered["Correlation_Coefficient"]
					p = df_filtered["adjusted_pvalues"]
					ci_low, ci_high = helper_funcs.correlations_format_val(df_filtered["CI_low"]), helper_funcs.correlations_format_val(df_filtered["CI_high"])
					r = helper_funcs.correlations_format_val(r, p)
					if "_" in r:
						r, sign = r.split("_")
					else:
						sign = ""
					table.cell(row_idx=inside_loop_ind, col_idx=outside_loop_ind).text = r
					table.cell(row_idx=inside_loop_ind, col_idx=outside_loop_ind+1).text = sign
					table.cell(row_idx=inside_loop_ind+1, col_idx=outside_loop_ind).text = "[" + ci_low + ", " + ci_high + "]"

	for cell in table.rows[0].cells:
		helper_funcs.set_cell_border(cell, top=global_vars.border_APA_word, bottom=global_vars.border_APA_word)
	for cell in table.rows[table_rows_len-1].cells:
		helper_funcs.set_cell_border(cell, bottom=global_vars.border_APA_word)

	for i in range(1, table_cols_len, 2):
		table.cell(row_idx=0, col_idx=i).merge(table.cell(row_idx=0, col_idx=i+1))
		table.cell(row_idx=0, col_idx=i).text = table.cell(row_idx=0, col_idx=i).text[:-1] # meging cells adds a paragraph break at the end; this gets rid of it
	for i in range(1, table_rows_len, 2):
		table.cell(row_idx=i, col_idx=0).merge(table.cell(row_idx=i+1, col_idx=0))
		table.cell(row_idx=i, col_idx=0).text = table.cell(row_idx=i, col_idx=0).text[:-1] # meging cells adds a paragraph break at the end; this gets rid of it
	for row in range(2, table_rows_len, 2):
		for col in range(1, table_cols_len, 2):
			table.cell(row_idx=row, col_idx=col).merge(table.cell(row_idx=row, col_idx=col+1))
			table.cell(row_idx=row, col_idx=col).alignment = WD_ALIGN_PARAGRAPH.CENTER
			helper_funcs.word_style(table.cell(row_idx=row, col_idx=col), size=9)

	for cell in table.rows[0].cells:
		helper_funcs.word_style(cell, italic=True)
		cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
		cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
	for cell in table.columns[0].cells:
		helper_funcs.word_style(cell, italic=True)
		cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

	for row in range(0, table_rows_len):
		for col in range(1, table_cols_len):
			cell = table.cell(row_idx=row, col_idx=col)
			if col >= 1:
				cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
			cell.paragraphs[0].paragraph_format.space_after = Inches(0)
			cell.paragraphs[0].paragraph_format.space_before = Inches(0)
			cell.paragraphs[0].paragraph_format.line_spacing = 1

	for row in range(1, table_rows_len, 2):
		for col in range(1, table_cols_len, 2):
			table.cell(row_idx=row, col_idx=col).paragraphs[0].paragraph_format.right_indent = Inches(-0.08)
			table.cell(row_idx=row, col_idx=col+1).paragraphs[0].paragraph_format.left_indent = Inches(-0.06)
			table.cell(row_idx=row, col_idx=col).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
			table.cell(row_idx=row, col_idx=col+1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

	doc = helper_funcs.set_autofit(doc)

	para = doc.add_paragraph("** ")
	para.add_run("p").italic = True
	para.add_run(" < 0.01")
	para.add_run().add_break()
	para.add_run("* ")
	para.add_run("p").italic = True
	para.add_run(" < {}".format(global_vars.alpha_threshold))
	doc.add_paragraph("Correlation coefficient used: {}".format(list(global_vars.master_dict.keys())[list(global_vars.master_dict.values()).index(global_vars.raw_corr_type)]))
	helper_funcs.add_correction_message_word(doc)

	helper_funcs.savefile(doc=doc)