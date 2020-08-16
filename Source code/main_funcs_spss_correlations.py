# MUltiple tests corrections and FOrmatted tables Software (MUFOS)
# Copyright (C) 2020  Nikolay Petrov, Vasil Atanasov, & Trevor Thompson
import global_vars
import helper_funcs

import pandas as pd
import numpy as np
from openpyxl import Workbook
from openpyxl.styles import Border, Side, Alignment, Font
from openpyxl.utils import get_column_letter
from openpyxl.utils.dataframe import dataframe_to_rows
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Inches

#-----------------------------------------------------------Modified raw data dataframe----------------------------------------------------
def spss_corr_generate_mod_raw_data_df(raw_data_df):
	if not(raw_data_df.columns[0] == "Correlations" or
		raw_data_df.iloc[1, 0] == "Spearman's rho" or raw_data_df.iloc[1, 0] == "Kendall's tau_b" or raw_data_df.iloc[1, 1] == "Pearson Correlation" or
		"Correlation is significant at the" in raw_data_df.iloc[raw_data_df.index[-1], 0] or
		raw_data_df.iloc[raw_data_df.index[-1], :][1:].isna().all() or
		raw_data_df.columns[1:].str.contains("Unnamed:").all()):
		
		raise Exception("The provided SPSS correlation table does not seem to be in the accepted format. Please use the \"How to export SPSS table?\" button for guidance or refer to the documentation.")

	mod_raw_data_df = raw_data_df.copy()

	if mod_raw_data_df.iloc[1, 1] == "Pearson Correlation":
		mod_raw_data_df.iloc[0, 0], mod_raw_data_df.iloc[0, 1] = "Variable", "Statistic"
		mod_raw_data_df = mod_raw_data_df.rename(columns=mod_raw_data_df.iloc[0]).drop(mod_raw_data_df.index[0]).reset_index(drop=True)
	elif mod_raw_data_df.iloc[1, 0] == "Kendall's tau_b" or mod_raw_data_df.iloc[1, 0] == "Spearman's rho":
		mod_raw_data_df.iloc[0, 1], mod_raw_data_df.iloc[0, 2] = "Variable", "Statistic"
		mod_raw_data_df = mod_raw_data_df.rename(columns=mod_raw_data_df.iloc[0]).drop(mod_raw_data_df.index[0]).reset_index(drop=True)
		correlation_label = mod_raw_data_df.iloc[0, 0]
		mod_raw_data_df["Statistic"].replace(to_replace="Correlation Coefficient", value=correlation_label, inplace=True)
		mod_raw_data_df.drop(columns = list(mod_raw_data_df.columns)[0], inplace=True)

	return mod_raw_data_df

#-----------------------------------------------------------Output dataframe----------------------------------------------------
def spss_corr_generate_output_df(mod_raw_data_df):
	'''
	Logic: 1) from the variable column, get the indices of all rows that have a variable 2) loop through those indices and create a standalone dataframe
	for each variable 3) from that dataframe, look at the correlations 4) find the diagonal (correlation is 1) 5) once the diagnoal is found
	then go through each subsequent column and get the needed stats stuff based on the kys of the output_dict
	'''
	correlation_label = mod_raw_data_df.iloc[0, 1]
	output_dict = {"var1": [], "var2": [], correlation_label: [], "pvalues": []}
	vars_row_ind_list = list(mod_raw_data_df[mod_raw_data_df["Variable"].notnull()].index)

	for ind, ele in enumerate(vars_row_ind_list):
		if ind != len(vars_row_ind_list)-1: # check for overflow
			current_df = mod_raw_data_df.iloc[vars_row_ind_list[ind]: vars_row_ind_list[ind+1]]
			row_corrs = current_df.iloc[0, 2:]
			diagonal_flag = False
			for var, corr in row_corrs.items():
				if corr == 1:
					diagonal_flag = True
					continue
				if diagonal_flag == True:
					output_dict["var1"].append(current_df.iloc[0, 0])
					output_dict["var2"].append(var)
					output_dict[correlation_label].append(corr)
					output_dict["pvalues"].append(current_df.loc[ele+1, var])

	output_df = pd.DataFrame(output_dict)

	return output_df

#-----------------------------------------------------------4. Saving data----------------------------------------------------
def spss_corr_apa_table_excel(mod_raw_data_df, output_df):
	# could very easily use the summ_corr_apa_table function here as I pass identical data - seperate is preferred, however, as might be updated/adjusted in the future
	correlation_label = mod_raw_data_df.iloc[0, 1]
	variables_list = list(mod_raw_data_df.columns)[2:]


	if global_vars.corr_table_triangle == "Upper triangle":
		header_cols = variables_list[1:]
		header_rows = variables_list[:-1]
	elif global_vars.corr_table_triangle == "Lower triangle":
		header_cols = variables_list[:-1]
		header_rows = variables_list[1:]
	elif global_vars.corr_table_triangle == "Both":
		header_cols, header_rows = variables_list, variables_list

	wb = Workbook()
	ws = wb.active

	ws.append([""] + header_cols)
	for ind,var in enumerate(header_rows):
		ws.cell(row=ind+2, column=1).value = var

	inside_loop_ind_start = 2
	for outside_loop_ind in range(2, len(header_rows) + 2):

		if global_vars.corr_table_triangle == "Upper triangle" or global_vars.corr_table_triangle == "Both":
			outside_loop_var = ws.cell(row=outside_loop_ind, column=1).value
		elif global_vars.corr_table_triangle == "Lower triangle":
			outside_loop_var = ws.cell(row=1, column=outside_loop_ind).value
		
		for inside_loop_ind in range(inside_loop_ind_start, len(header_cols) + 2):
			
			if global_vars.corr_table_triangle == "Upper triangle" or global_vars.corr_table_triangle == "Both":
				inside_loop_var = ws.cell(row=1, column=inside_loop_ind).value
			elif global_vars.corr_table_triangle == "Lower triangle":
				inside_loop_var = ws.cell(row=inside_loop_ind, column=1).value

			if outside_loop_var == inside_loop_var:
				ws.cell(row=outside_loop_ind, column=inside_loop_ind).value = 1
			else:
				# here query method is not preferred as it is not only slower as it is much smaller dataset but also cannot refer to two different vaiables (colname and val)
				df_filtered = output_df[((output_df["var1"]==outside_loop_var) & (output_df["var2"]==inside_loop_var)) | ((output_df["var1"]==inside_loop_var) & (output_df["var2"]==outside_loop_var))].iloc[0]
				r = df_filtered[correlation_label]
				p = df_filtered["adjusted_pvalues"]
				r = helper_funcs.correlations_format_val(r, p)

				if global_vars.corr_table_triangle == "Upper triangle" or global_vars.corr_table_triangle == "Both":
					ws.cell(row=outside_loop_ind, column=inside_loop_ind).value = r
				elif global_vars.corr_table_triangle == "Lower triangle":
					ws.cell(row=inside_loop_ind, column=outside_loop_ind).value = r
		
		if global_vars.corr_table_triangle != "Both":
			inside_loop_ind_start += 1


	for row in range(1, len(header_rows) + 2):
		for cell in ws[row]:
			cell.alignment = global_vars.alignment_center

	for cell in ws[1]:
		cell.font = global_vars.font_header

	for row in range(2, len(header_rows) + 2):
		ws.cell(row=row, column=1).font = global_vars.font_header

	for cell in ws[len(header_rows) + 1]:
		cell.border = Border(bottom=global_vars.border_APA)

	for cell in ws[1]:
		cell.border = Border(top=global_vars.border_APA, bottom=global_vars.border_APA)

	table_notes = ["**p < 0.01", "*p < {}".format(global_vars.alpha_threshold)]
	table_notes.append("Correlation Coefficient used: {}".format(correlation_label))
	helper_funcs.add_table_notes(ws, table_notes)

	helper_funcs.savefile(wb=wb)

def spss_corr_apa_table_word(mod_raw_data_df, output_df):
	# could very easily use the summ_corr_apa_table function here as I pass identical data - seperate is preferred, however, as might be updated/adjusted in the future
	correlation_label = mod_raw_data_df.iloc[0, 1]
	variables_list = list(mod_raw_data_df.columns)[2:]

	if global_vars.corr_table_triangle == "Upper triangle":
		header_cols = variables_list[1:]
		header_rows = variables_list[:-1]
	elif global_vars.corr_table_triangle == "Lower triangle":
		header_cols = variables_list[:-1]
		header_rows = variables_list[1:]
	elif global_vars.corr_table_triangle == "Both":
		header_cols = [x for x in variables_list] # not straight up variables_list; pass by value vs pass by reference 
		header_rows = [x for x in variables_list]
	# this adds an empty column where the significance signs will be placed for better presentation
	# the code looks slightly wrong as it return a None array but it works
	[header_cols.insert(x, "") for x in range(1, len(header_cols)*2, 2)]

	doc = Document()
	table_rows_len = len(header_rows) + 1
	table_cols_len = len(header_cols) + 1
	table = doc.add_table(rows=table_rows_len, cols=table_cols_len)

	for ind, var in enumerate([""] + header_cols):
		table.cell(row_idx=0, col_idx=ind).text = var
	for ind, var in enumerate([""] + header_rows):
		table.cell(row_idx=ind, col_idx=0).text = var


	if global_vars.corr_table_triangle == "Upper triangle":
		inside_loop_ind_start = 1
		for outside_loop_ind in range(1, table_rows_len):
			outside_loop_var = table.cell(row_idx=outside_loop_ind, col_idx=0).text
			for inside_loop_ind in range(inside_loop_ind_start, table_cols_len):
				inside_loop_var = table.cell(row_idx=0, col_idx=inside_loop_ind).text
				if inside_loop_var == "" or outside_loop_var == "": # this allows to skip the columns designed for the significance signs
					continue
				else:
					# here query method is not preferred as it is not only slower as it is much smaller dataset but also cannot refer to two different vaiables (colname and val)
					df_filtered = output_df[((output_df["var1"]==outside_loop_var) & (output_df["var2"]==inside_loop_var)) | ((output_df["var1"]==inside_loop_var) & (output_df["var2"]==outside_loop_var))].iloc[0]
					r = df_filtered[correlation_label]
					p = df_filtered["adjusted_pvalues"]
					r = helper_funcs.correlations_format_val(r, p)
					if "_" in r:
						r, sign = r.split("_")
					else:
						sign = ""
					table.cell(row_idx=outside_loop_ind, col_idx=inside_loop_ind).text = r
					table.cell(row_idx=outside_loop_ind, col_idx=inside_loop_ind+1).text = sign
			inside_loop_ind_start += 2
	elif global_vars.corr_table_triangle == "Lower triangle":
		inside_loop_ind_start = 1
		for outside_loop_ind in range(1, table_cols_len):
			outside_loop_var = table.cell(row_idx=0, col_idx=outside_loop_ind).text
			if outside_loop_var == "":
				continue
			else:
				for inside_loop_ind in range(inside_loop_ind_start, table_rows_len):
					inside_loop_var = table.cell(row_idx=inside_loop_ind, col_idx=0).text
					# here query method is not preferred as it is not only slower as it is much smaller dataset but also cannot refer to two different vaiables (colname and val)
					df_filtered = output_df[((output_df["var1"]==outside_loop_var) & (output_df["var2"]==inside_loop_var)) | ((output_df["var1"]==inside_loop_var) & (output_df["var2"]==outside_loop_var))].iloc[0]
					r = df_filtered[correlation_label]
					p = df_filtered["adjusted_pvalues"]
					r = helper_funcs.correlations_format_val(r, p)
					if "_" in r:
						r, sign = r.split("_")
					else:
						sign = ""
					table.cell(row_idx=inside_loop_ind, col_idx=outside_loop_ind).text = r
					table.cell(row_idx=inside_loop_ind, col_idx=outside_loop_ind+1).text = sign
				inside_loop_ind_start += 1
	elif global_vars.corr_table_triangle == "Both":
		inside_loop_ind_start = 1
		for outside_loop_ind in range(1, table_cols_len):
			outside_loop_var = table.cell(row_idx=0, col_idx=outside_loop_ind).text
			if outside_loop_var == "":
				continue
			else:
				for inside_loop_ind in range(inside_loop_ind_start, table_rows_len):
					inside_loop_var = table.cell(row_idx=inside_loop_ind, col_idx=0).text
					if inside_loop_var == "": # this allows to skip the columns designed for the significance signs
						continue
					else:
						if outside_loop_var == inside_loop_var:
							table.cell(row_idx=inside_loop_ind, col_idx=outside_loop_ind).text = "1"
						else:
						# here query method is not preferred as it is not only slower as it is much smaller dataset but also cannot refer to two different vaiables (colname and val)
							df_filtered = output_df[((output_df["var1"]==outside_loop_var) & (output_df["var2"]==inside_loop_var)) | ((output_df["var1"]==inside_loop_var) & (output_df["var2"]==outside_loop_var))].iloc[0]
							r = df_filtered[correlation_label]
							p = df_filtered["adjusted_pvalues"]
							r = helper_funcs.correlations_format_val(r, p)
							if "_" in r:
								r, sign = r.split("_")
							else:
								sign = ""
							table.cell(row_idx=inside_loop_ind, col_idx=outside_loop_ind).text = r
							table.cell(row_idx=inside_loop_ind, col_idx=outside_loop_ind+1).text = sign


	for cell in table.rows[0].cells:
		helper_funcs.set_cell_border(cell, top=global_vars.border_APA_word, bottom=global_vars.border_APA_word)
	for cell in table.rows[table_rows_len-1].cells:
		helper_funcs.set_cell_border(cell, bottom=global_vars.border_APA_word)

	for cell in table.rows[0].cells:
		helper_funcs.word_style(cell, italic=True)
	for cell in table.columns[0].cells:
		helper_funcs.word_style(cell, italic=True)

	for i in range(1, table_cols_len, 2):
		table.cell(row_idx=0, col_idx=i).merge(table.cell(row_idx=0, col_idx=i+1))
		table.cell(row_idx=0, col_idx=i).text = table.cell(row_idx=0, col_idx=i).text[:-1] # meging cells adds a paragraph break at the end; this gets rid of it

	for cell in table.rows[0].cells:
		helper_funcs.word_style(cell, italic=True)
		cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
	for cell in table.columns[0].cells:
		helper_funcs.word_style(cell, italic=True)
		cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

	for row in range(0, table_rows_len):
		for cell in table.rows[row].cells:
			cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
			cell.paragraphs[0].paragraph_format.space_after = Inches(0)
			cell.paragraphs[0].paragraph_format.space_before = Inches(0)
			cell.paragraphs[0].paragraph_format.line_spacing = 1

	for row in range(1, table_rows_len):
		for col in range(1, table_cols_len, 2):
			table.cell(row_idx=row, col_idx=col).paragraphs[0].paragraph_format.right_indent = Inches(-0.08)
			table.cell(row_idx=row, col_idx=col+1).paragraphs[0].paragraph_format.left_indent = Inches(-0.06)
			table.cell(row_idx=row, col_idx=col).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
			table.cell(row_idx=row, col_idx=col+1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

	doc = helper_funcs.set_autofit(doc)

	para = doc.add_paragraph("** ")
	para.add_run("p").italic = True
	para.add_run(" < 0.01")
	para.add_run().add_break()
	para.add_run("* ")
	para.add_run("p").italic = True
	para.add_run(" < {}".format(global_vars.alpha_threshold))
	doc.add_paragraph("Correlation Coefficient used: {}".format(correlation_label))
	helper_funcs.add_correction_message_word(doc)

	helper_funcs.savefile(doc=doc)
