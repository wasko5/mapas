# MUltiple tests corrections and FOrmatted tables Software (MUFOS)
# Copyright (C) 2020  Nikolay Petrov, Vasil Atanasov, & Trevor Thompson
import global_vars
import helper_funcs

import pandas as pd
import numpy as np
from openpyxl import Workbook
from openpyxl.styles import Border, Side, Alignment, Font
from openpyxl.utils import get_column_letter
from openpyxl.utils.dataframe import dataframe_to_rows
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

#-----------------------------------------------------------Modified raw data dataframe----------------------------------------------------
def spss_indttest_generate_mod_raw_data_df(raw_data_df):
	mod_raw_data_df = raw_data_df.copy()

	# renaming columns as the multi-level formatting of the spss table makes it difficult to work with
	renaming_dict = {}
	new_col_names = list(mod_raw_data_df.columns[:2]) + list(mod_raw_data_df.iloc[1][2:-1]) + [mod_raw_data_df.columns[-1]]
	for i in range(0, len(mod_raw_data_df.columns)):
		renaming_dict[mod_raw_data_df.columns[i]] = new_col_names[i]

	mod_raw_data_df = mod_raw_data_df.rename(columns=renaming_dict).drop([mod_raw_data_df.index[0], mod_raw_data_df.index[1]]).reset_index(drop=True)
	
	if mod_raw_data_df.columns[0] != "Independent Samples Test":
		raise Exception("The SPSS table provided does not seem to be one from Independent Samples t-test. Please try with another file or refer to the documentation for help.")

	fields_to_check = ["F", "Sig.", "t", "df", "Sig. (2-tailed)"]
	for field in fields_to_check:
		if field not in mod_raw_data_df.columns:
			raise Exception("The column \'{}\' was not found in the file. Please try entering a correct data file. Please see the documentation for help.".format(field))

	return mod_raw_data_df

#-----------------------------------------------------------Output dataframe----------------------------------------------------
def spss_indttest_generate_output_df(mod_raw_data_df):
	output_dict = {"Variable": [], "All_Mean": [], "All_SD": [], global_vars.spss_indttest_groupOneLabel+"_Mean": [], global_vars.spss_indttest_groupOneLabel+"_SD": [],
					global_vars.spss_indttest_groupTwoLabel+"_Mean": [], global_vars.spss_indttest_groupTwoLabel+"_SD": [], "Degrees of Freedom": [], "t": [],
					global_vars.effect_size_choice: [], "pvalues": []}

	for row in range(1, len(mod_raw_data_df), 2):
		levene_p = mod_raw_data_df["Sig."][row]
		if levene_p > 0.05:
			row_to_read = row
		else:
			row_to_read = row+1
		t = mod_raw_data_df["t"][row_to_read]
		effect_size = helper_funcs.calc_ttest_effect_size(effect_size_choice = global_vars.effect_size_choice, t=t, n1=global_vars.spss_indttest_nOne, n2=global_vars.spss_indttest_nTwo)

		output_dict["Variable"].append(mod_raw_data_df["Independent Samples Test"][row])
		output_dict["Degrees of Freedom"].append(mod_raw_data_df["df"][row_to_read])
		output_dict["t"].append(t)
		output_dict[global_vars.effect_size_choice].append(effect_size)
		output_dict["pvalues"].append(mod_raw_data_df["Sig. (2-tailed)"][row_to_read])
		
	output_dict["All_Mean"] = [""] * len(output_dict["Variable"])
	output_dict["All_SD"] = [""] * len(output_dict["Variable"])
	output_dict[global_vars.spss_indttest_groupOneLabel + "_Mean"] = [""] * len(output_dict["Variable"])
	output_dict[global_vars.spss_indttest_groupOneLabel + "_SD"] = [""] * len(output_dict["Variable"])
	output_dict[global_vars.spss_indttest_groupTwoLabel + "_Mean"] = [""] * len(output_dict["Variable"])
	output_dict[global_vars.spss_indttest_groupTwoLabel + "_SD"] = [""] * len(output_dict["Variable"])

	output_df = pd.DataFrame(output_dict)

	return output_df

#-----------------------------------------------------------Saving data----------------------------------------------------
def spss_indttest_apa_table_excel(mod_raw_data_df, output_df):
	output_df.drop(columns=["pvalues"], inplace=True)

	pd.options.mode.chained_assignment = None
	output_df[list(output_df.columns)[7:-1]] = output_df[list(output_df.columns)[7:-1]].applymap(lambda x: "{:.2f}".format(x))

	output_df["adjusted_pvalues"] = output_df["adjusted_pvalues"].map(helper_funcs.pvalue_formatting)
	pd.options.mode.chained_assignment = "warn"

	wb = Workbook()
	ws = wb.active
	
	ws.cell(row=1, column=1).value = "Variable"
	ws.merge_cells('A1:A2')
	ws.cell(row=1, column=1).font = global_vars.font_header
	
	ws.cell(row=1, column=2).value = "All, n=?"
	ws.merge_cells('B1:C1')
	
	if global_vars.spss_indttest_nOne != -1:
		n1_label = global_vars.spss_indttest_nOne
	else:
		n1_label = "?"
	ws.cell(row=1, column=4).value = "{g}, n={n}".format(g=global_vars.spss_indttest_groupOneLabel, n=n1_label)
	ws.merge_cells('D1:E1')
	
	if global_vars.spss_indttest_nOne != -1:
		n2_label = global_vars.spss_indttest_nTwo
	else:
		n2_label = "?"
	ws.cell(row=1, column=6).value = "{g}, n={n}".format(g=global_vars.spss_indttest_groupTwoLabel, n=n2_label)
	ws.merge_cells('F1:G1')
	
	ws.cell(row=1, column=8).value = "df"
	ws.merge_cells('H1:H2')
	ws.cell(row=1, column=8).font = global_vars.font_header
	
	ws.cell(row=1, column=9).value = "t"
	ws.merge_cells('I1:I2')
	ws.cell(row=1, column=9).font = global_vars.font_header
	
	ws.cell(row=1, column=10).value = global_vars.effect_size_choice
	ws.merge_cells('J1:J2')
	ws.cell(row=1, column=10).font = global_vars.font_header
	
	ws.cell(row=1, column=11).value = "p"
	ws.merge_cells('K1:K2')
	ws.cell(row=1, column=11).font = global_vars.font_header

	for col in range(2,7,2):
		ws.cell(row=2, column=col).value = "M"
		ws.cell(row=2, column=col).font = global_vars.font_header
		ws.cell(row=2, column=col+1).value = "SD"
		ws.cell(row=2, column=col+1).font = global_vars.font_header

	for row in dataframe_to_rows(output_df, index=False, header=False):
		ws.append(row)

	for row in range(1, len(output_df)+3):
		for cell in ws[row]:
			cell.alignment = global_vars.alignment_center

	for cell in ws[1]:
		cell.border = Border(top=global_vars.border_APA)
	for cell in ws[2] + ws[len(output_df)+2]:
		cell.border = Border(bottom=global_vars.border_APA)

	if global_vars.effect_size_choice == "None":
		ws.delete_cols(10)

	table_notes = ["Means and Standard Deviations cannot be read from the SPSS table. Please add them yourself or remove those columns if they are not needed."]
	helper_funcs.add_table_notes(ws, table_notes)

	helper_funcs.savefile(wb=wb)

def spss_indttest_apa_table_word(mod_raw_data_df, output_df):
	# very similar to raw_corr_indttest_apa_table func but separate as might be udpated/adjusted in the future
	output_df.drop(columns=["pvalues"], inplace=True)

	pd.options.mode.chained_assignment = None
	output_df[list(output_df.columns)[7:-1]] = output_df[list(output_df.columns)[7:-1]].applymap(lambda x: "{:.2f}".format(x))

	output_df["adjusted_pvalues"] = output_df["adjusted_pvalues"].map(helper_funcs.pvalue_formatting)
	pd.options.mode.chained_assignment = "warn"

	doc = Document()
	table_rows_len = len(output_df) + 2
	table_cols_len = len(output_df.columns)
	table = doc.add_table(rows=table_rows_len, cols=table_cols_len)

	table.cell(row_idx=0, col_idx=0).text = "Variable"
	table.cell(row_idx=0, col_idx=0).merge(table.cell(row_idx=1, col_idx=0))
	helper_funcs.word_style(table.cell(row_idx=0, col_idx=0), italic=True)

	table.cell(row_idx=0, col_idx=1).text = "All, n=?"
	table.cell(row_idx=0, col_idx=1).merge(table.cell(row_idx=0, col_idx=2))

	if global_vars.spss_indttest_nOne != -1:
		n1_label = global_vars.spss_indttest_nOne
	else:
		n1_label = "?"
	table.cell(row_idx=0, col_idx=3).text = "{g}, n={n}".format(g=global_vars.spss_indttest_groupOneLabel, n=n1_label)
	table.cell(row_idx=0, col_idx=3).merge(table.cell(row_idx=0, col_idx=4))

	if global_vars.spss_indttest_nOne != -1:
		n2_label = global_vars.spss_indttest_nTwo
	else:
		n2_label = "?"
	table.cell(row_idx=0, col_idx=5).text = "{g}, n={n}".format(g=global_vars.spss_indttest_groupTwoLabel, n=n2_label)
	table.cell(row_idx=0, col_idx=5).merge(table.cell(row_idx=0, col_idx=6))

	table.cell(row_idx=0, col_idx=7).text = "df"
	table.cell(row_idx=0, col_idx=7).merge(table.cell(row_idx=1, col_idx=7))
	helper_funcs.word_style(table.cell(row_idx=0, col_idx=7), italic=True)

	table.cell(row_idx=0, col_idx=8).text = "t"
	table.cell(row_idx=0, col_idx=8).merge(table.cell(row_idx=1, col_idx=8))
	helper_funcs.word_style(table.cell(row_idx=0, col_idx=8), italic=True)

	table.cell(row_idx=0, col_idx=9).text = global_vars.effect_size_choice
	if global_vars.effect_size_choice != "None": # otherwise cant remove with delete columns below if merged; see helper_funcs.delete_columns_word func notes
		table.cell(row_idx=0, col_idx=9).merge(table.cell(row_idx=1, col_idx=9))
	helper_funcs.word_style(table.cell(row_idx=0, col_idx=9), italic=True)
	
	table.cell(row_idx=0, col_idx=10).text = "p"
	table.cell(row_idx=0, col_idx=10).merge(table.cell(row_idx=1, col_idx=10))
	helper_funcs.word_style(table.cell(row_idx=0, col_idx=10), italic=True)

	for col in range(1, 6, 2):
		table.cell(row_idx=1, col_idx=col).text = "M"
		helper_funcs.word_style(table.cell(row_idx=1, col_idx=col), italic=True)
		table.cell(row_idx=1, col_idx=col+1).text = "SD"
		helper_funcs.word_style(table.cell(row_idx=1, col_idx=col+1), italic=True)

	for row in range(2, table_rows_len):
		for col in range(0, table_cols_len):
			table.cell(row_idx=row, col_idx=col).text = output_df.iloc[row-2, col]

	for row in range(0, table_rows_len):
		for cell in table.rows[row].cells:
			cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
			cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

	for cell in table.rows[0].cells:
		helper_funcs.set_cell_border(cell, top=global_vars.border_APA_word)
	for cell in table.rows[2].cells:
		helper_funcs.set_cell_border(cell, top=global_vars.border_APA_word)
	for cell in table.rows[table_rows_len-1].cells:
		helper_funcs.set_cell_border(cell, bottom=global_vars.border_APA_word)

	if global_vars.effect_size_choice == "None":
		helper_funcs.delete_columns_word(table, [9])

	doc = helper_funcs.set_autofit(doc)

	doc.add_paragraph("Means and Standard Deviations cannot be read from the SPSS table. Please add them yourself or remove those columns if they are not needed.")
	helper_funcs.add_correction_message_word(doc)

	helper_funcs.savefile(doc=doc)