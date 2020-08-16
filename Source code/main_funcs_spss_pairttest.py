# MUltiple tests corrections and FOrmatted tables Software (MUFOS)
# Copyright (C) 2020  Nikolay Petrov, Vasil Atanasov, & Trevor Thompson
import global_vars
import helper_funcs

import pandas as pd
import numpy as np
from openpyxl import Workbook
from openpyxl.styles import Border, Side, Alignment, Font
from openpyxl.utils import get_column_letter
from openpyxl.utils.dataframe import dataframe_to_rows
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

#-----------------------------------------------------------Modified raw data dataframe----------------------------------------------------
def spss_pairttest_generate_mod_raw_data_df(raw_data_df):
	if (len(raw_data_df.columns) != 10 or raw_data_df.iloc[1,2] != "Mean" or raw_data_df.iloc[1,3] != "Std. Deviation"
			or raw_data_df.iloc[1,4] != "Std. Error Mean" or raw_data_df.iloc[1,5] != "95% Confidence Interval of the Difference"
			or raw_data_df.iloc[2,6] != "Upper" or raw_data_df.iloc[0,7] != "t" or raw_data_df.iloc[0,8] != "df"
			or raw_data_df.iloc[0,9] != "Sig. (2-tailed)" or len(raw_data_df) <= 3):
		raise Exception("The data provided could not be read correctly. Please see the documentation for help.")
		
	mod_raw_data_df = raw_data_df.copy()
	renaming_dict = {}
	new_col_names = [mod_raw_data_df.columns[0]] + ["Variables", "Mean difference", "Std difference", "Std error difference"] + [mod_raw_data_df.iloc[1, 5]] + ["Upper", "t", "df", "p"]
	for i in range(0, len(mod_raw_data_df.columns)):
		renaming_dict[mod_raw_data_df.columns[i]] = new_col_names[i]

	mod_raw_data_df = mod_raw_data_df.rename(columns=renaming_dict).drop([mod_raw_data_df.index[0], mod_raw_data_df.index[1], mod_raw_data_df.index[2]]).reset_index(drop=True)

	return mod_raw_data_df

#-----------------------------------------------------------Output dataframe----------------------------------------------------
def spss_pairttest_generate_output_df(mod_raw_data_df):
	output_dict={}
	t_stat_list = list(mod_raw_data_df["t"])
	effectsize_list = []

	for t in t_stat_list:
		effect_size = helper_funcs.calc_ttest_effect_size(global_vars.effect_size_choice, t=t, n1=global_vars.spss_pairttest_n, n2=global_vars.spss_pairttest_n)
		effectsize_list.append(effect_size)
			
	output_dict["Variables"] = list(mod_raw_data_df["Variables"])
	output_dict["Time 1_Mean"] = [""] * len(output_dict["Variables"])
	output_dict["Time 1_SD"] = [""] * len(output_dict["Variables"])
	output_dict["Time 2_Mean"] = [""] * len(output_dict["Variables"])
	output_dict["Time 2_SD"] = [""] * len(output_dict["Variables"])
	output_dict["df"] = list(mod_raw_data_df["df"])
	output_dict["t"] = t_stat_list
	output_dict[global_vars.effect_size_choice] = effectsize_list
	output_dict["pvalues"] = list(mod_raw_data_df["p"])

	output_df = pd.DataFrame(output_dict)
	
	return output_df

#-----------------------------------------------------------Saving data----------------------------------------------------
def spss_pairttest_apa_table_excel(mod_raw_data_df, output_df):
	output_df.drop(columns=["pvalues"], inplace=True)

	pd.options.mode.chained_assignment = None
	output_df[list(output_df.columns)[5:-1]] = output_df[list(output_df.columns)[5:-1]].applymap(lambda x: "{:.2f}".format(x))

	output_df["adjusted_pvalues"] = output_df["adjusted_pvalues"].map(helper_funcs.pvalue_formatting)
	pd.options.mode.chained_assignment = "warn"

	wb = Workbook()
	ws = wb.active
	
	ws.cell(row=1, column=1).value = "Variable"
	ws.merge_cells('A1:A2')
	ws.cell(row=1, column=1).font = global_vars.font_header
	
	ws.cell(row=1, column=2).value = "Time 1"
	ws.merge_cells('B1:C1')
	
	ws.cell(row=1, column=4).value = "Time 2"
	ws.merge_cells('D1:E1')
	
	ws.cell(row=1, column=6).value = "df"
	ws.merge_cells('F1:F2')
	ws.cell(row=1, column=6).font = global_vars.font_header
	
	ws.cell(row=1, column=7).value = "t"
	ws.merge_cells('G1:G2')
	ws.cell(row=1, column=7).font = global_vars.font_header
	
	ws.cell(row=1, column=8).value = global_vars.effect_size_choice
	ws.merge_cells('H1:H2')
	ws.cell(row=1, column=8).font = global_vars.font_header
	
	ws.cell(row=1, column=9).value = "p"
	ws.merge_cells('I1:I2')
	ws.cell(row=1, column=9).font = global_vars.font_header
	
	for col in range(2,5,2):
		ws.cell(row=2, column=col).value = "M"
		ws.cell(row=2, column=col).font = global_vars.font_header
		ws.cell(row=2, column=col+1).value = "SD"
		ws.cell(row=2, column=col+1).font = global_vars.font_header
	
	for row in dataframe_to_rows(output_df, index=False, header=False):
		ws.append(row)

	for row in range(1, len(output_df)+3):
		for cell in ws[row]:
			cell.alignment = global_vars.alignment_center
	
	for cell in ws[1]:
		cell.border = Border(top=global_vars.border_APA)
	for cell in ws[2] + ws[len(output_df)+2]:
		cell.border = Border(bottom=global_vars.border_APA)

	if global_vars.effect_size_choice == "None":
		ws.delete_cols(8)

	table_notes = ["Means and Standard Deviations cannot be read from the SPSS table. Please add them yourself or remove those columns if they are not needed."]
	helper_funcs.add_table_notes(ws, table_notes)

	helper_funcs.savefile(wb=wb)

def spss_pairttest_apa_table_word(mod_raw_data_df, output_df):
	output_df.drop(columns=["pvalues"], inplace=True)

	pd.options.mode.chained_assignment = None
	output_df[list(output_df.columns)[5:-1]] = output_df[list(output_df.columns)[5:-1]].applymap(lambda x: "{:.2f}".format(x))

	output_df["adjusted_pvalues"] = output_df["adjusted_pvalues"].map(helper_funcs.pvalue_formatting)
	pd.options.mode.chained_assignment = "warn"

	doc = Document()
	table_rows_len = len(output_df) + 2
	table_cols_len = len(output_df.columns)
	table = doc.add_table(rows=table_rows_len, cols=table_cols_len)

	table.cell(row_idx=0, col_idx=0).text = "Variable"
	table.cell(row_idx=0, col_idx=0).merge(table.cell(row_idx=1, col_idx=0))
	helper_funcs.word_style(table.cell(row_idx=0, col_idx=0), italic=True)

	table.cell(row_idx=0, col_idx=1).text = "Time 1"
	table.cell(row_idx=0, col_idx=1).merge(table.cell(row_idx=0, col_idx=2))
	
	table.cell(row_idx=0, col_idx=3).text = "Time 2"
	table.cell(row_idx=0, col_idx=3).merge(table.cell(row_idx=0, col_idx=4))

	table.cell(row_idx=0, col_idx=5).text = "df"
	table.cell(row_idx=0, col_idx=5).merge(table.cell(row_idx=1, col_idx=5))
	helper_funcs.word_style(table.cell(row_idx=0, col_idx=5), italic=True)
	
	table.cell(row_idx=0, col_idx=6).text = "t"
	table.cell(row_idx=0, col_idx=6).merge(table.cell(row_idx=1, col_idx=6))
	helper_funcs.word_style(table.cell(row_idx=0, col_idx=6), italic=True)
	
	table.cell(row_idx=0, col_idx=7).text = global_vars.effect_size_choice
	if global_vars.effect_size_choice != "None": # otherwise cant remove with delete columns below if merged; see helper_funcs.delete_columns_word func notes
		table.cell(row_idx=0, col_idx=7).merge(table.cell(row_idx=1, col_idx=7))
	helper_funcs.word_style(table.cell(row_idx=0, col_idx=7), italic=True)
	
	table.cell(row_idx=0, col_idx=8).text = "p"
	table.cell(row_idx=0, col_idx=8).merge(table.cell(row_idx=1, col_idx=8))
	helper_funcs.word_style(table.cell(row_idx=0, col_idx=8), italic=True)

	for col in range(1, 4, 2):
		table.cell(row_idx=1, col_idx=col).text = "M"
		helper_funcs.word_style(table.cell(row_idx=1, col_idx=col), italic=True)
		table.cell(row_idx=1, col_idx=col+1).text = "SD"
		helper_funcs.word_style(table.cell(row_idx=1, col_idx=col+1), italic=True)
	
	for row in range(2, table_rows_len):
		for col in range(0, table_cols_len):
			table.cell(row_idx=row, col_idx=col).text = output_df.iloc[row-2, col]

	for row in range(0, table_rows_len):
		for cell in table.rows[row].cells:
			cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
			cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

	for cell in table.rows[0].cells:
		helper_funcs.set_cell_border(cell, top=global_vars.border_APA_word)
	for cell in table.rows[2].cells:
		helper_funcs.set_cell_border(cell, top=global_vars.border_APA_word)
	for cell in table.rows[table_rows_len-1].cells:
		helper_funcs.set_cell_border(cell, bottom=global_vars.border_APA_word)

	if global_vars.effect_size_choice == "None":
		helper_funcs.delete_columns_word(table, [7])

	doc = helper_funcs.set_autofit(doc)

	doc.add_paragraph("Means and Standard Deviations cannot be read from the SPSS table. Please add them yourself or remove those columns if they are not needed.")
	helper_funcs.add_correction_message_word(doc)

	helper_funcs.savefile(doc=doc)
