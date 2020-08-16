# MUltiple tests corrections and FOrmatted tables Software (MUFOS)
# Copyright (C) 2020  Nikolay Petrov, Vasil Atanasov, & Trevor Thompson
import global_vars
import helper_funcs

import pandas as pd
import numpy as np
from scipy import stats
import statsmodels.api as sm
from openpyxl import Workbook
from openpyxl.styles import Border, Side, Alignment, Font
from openpyxl.utils import get_column_letter
from openpyxl.utils.dataframe import dataframe_to_rows
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

#-----------------------------------------------------------Output dataframe----------------------------------------------------
def raw_mr_generate_output_df(mod_raw_data_df):	
	formula = global_vars.raw_mr_outcomevar + "~"
	for var in global_vars.raw_mr_predictors:
		formula = formula + var + "+"
	formula = formula[:-1] # the loop leaves a trailing "+" sign so this gets rid of it

	result = sm.OLS.from_formula(formula=formula, data=mod_raw_data_df).fit()
	
	# standardization of the raw data dataframe so it can also produce beta values for the output
	mod_raw_data_df_z = mod_raw_data_df.apply(lambda x: stats.zscore(x, nan_policy="omit"))
	result_z = sm.OLS.from_formula(formula=formula, data=mod_raw_data_df_z).fit()
	beta_list = list(result_z.params)

	CI_B_list = []
	for index, row in result.conf_int().iterrows(): # weirdly enough result_conf_int() returns a dataframe with the CIs
		CI_B_list.append(list(row))
		
	CI_beta_list = []
	for index, row in result_z.conf_int().iterrows():
		CI_beta_list.append(list(row))

	output_dict = {}
	output_dict["Variable"] = ["(Constant)"] + global_vars.raw_mr_predictors
	output_dict["B"] = list(result.params)
	output_dict["Std Err B"] = list(result.bse)
	output_dict["95% CI B"] = ["[" + ", ".join("{:.2f}".format(e) for e in pair) + "]" for pair in CI_B_list]
	output_dict["beta"] = beta_list
	output_dict["Std Err beta"] = list(result_z.bse) # might not need - if so, amend the apa_table func
	output_dict["95% CI beta"] = ["[" + ", ".join("{:.2f}".format(e) for e in pair) + "]" for pair in CI_beta_list] # might not need
	output_dict["t"] = list(result.tvalues)
	output_dict["pvalues"] = list(result.pvalues)
	output_dict["R2"] = [result.rsquared] + [np.nan] * (len(output_dict["Variable"])-1)
	output_dict["R2adj"] = [result.rsquared_adj] + [np.nan] * (len(output_dict["Variable"])-1)

	output_df = pd.DataFrame(output_dict)
	
	return output_df

#-----------------------------------------------------------Saving data----------------------------------------------------
def raw_mr_apa_table_excel(mod_raw_data_df, output_df):
	mod_output_df = output_df[["Variable","B","95% CI B","beta","t","adjusted_pvalues"]]

	pd.options.mode.chained_assignment = None
	mod_output_df[["B", "beta", "t"]] = mod_output_df[["B", "beta", "t"]].applymap(lambda x: "{:.2f}".format(x))

	mod_output_df["adjusted_pvalues"] = mod_output_df["adjusted_pvalues"].map(helper_funcs.pvalue_formatting)

	mod_output_df.rename(columns = {"95% CI B": "95% CI", "adjusted_pvalues": "p"}, inplace=True)
	mod_output_df.loc[0, "beta"] = "" # removes the beta value for constant as it is always 0

	wb = Workbook()
	ws = wb.active

	ws.append(list(mod_output_df.columns))

	for row in dataframe_to_rows(mod_output_df, index=False, header=False):
		ws.append(row)

	for cell in ws[1]:
		cell.font = global_vars.font_header

	for cell in ws[1] + ws[len(mod_output_df)+1]:
		cell.border = Border(top=global_vars.border_APA, bottom=global_vars.border_APA)
	for cell in ws[len(mod_output_df)+1]:
		cell.border = Border(bottom=global_vars.border_APA)

	for row in range(1, len(mod_output_df)+2):
		for cell in ws[row]:
			cell.alignment = global_vars.alignment_center

	table_notes = ["R squared adjusted = {R}".format(R="{:.2f}".format(output_df["R2adj"][0])), "Dependent Variable: {DV}".format(DV=global_vars.raw_mr_outcomevar)]
	helper_funcs.add_table_notes(ws, table_notes)

	helper_funcs.savefile(wb=wb)

def raw_mr_apa_table_word(mod_raw_data_df, output_df):
	mod_output_df = output_df[["Variable","B","95% CI B","beta","t","adjusted_pvalues"]]

	pd.options.mode.chained_assignment = None
	mod_output_df[["B", "beta", "t"]] = mod_output_df[["B", "beta", "t"]].applymap(lambda x: "{:.2f}".format(x))

	mod_output_df["adjusted_pvalues"] = mod_output_df["adjusted_pvalues"].map(helper_funcs.pvalue_formatting)

	mod_output_df.rename(columns = {"95% CI B": "95% CI", "adjusted_pvalues": "p"}, inplace=True)
	mod_output_df.loc[0, "beta"] = "" # removes the beta value for constant as it is always 0

	doc = Document()
	table_rows_len = len(mod_output_df) + 1
	table_cols_len = len(mod_output_df.columns)
	table = doc.add_table(rows=table_rows_len, cols=table_cols_len)

	for ind, var in enumerate(mod_output_df.columns):
		table.cell(row_idx=0, col_idx=ind).text = var

	for row in range(1, table_rows_len):
		for col in range(0, table_cols_len):
			table.cell(row_idx=row, col_idx=col).text = mod_output_df.iloc[row-1, col]

	for cell in table.rows[0].cells:
		helper_funcs.word_style(cell, italic=True)

	for cell in table.rows[0].cells:
		helper_funcs.set_cell_border(cell, top=global_vars.border_APA_word, bottom=global_vars.border_APA_word)
	for cell in table.rows[table_rows_len-1].cells:
		helper_funcs.set_cell_border(cell, bottom=global_vars.border_APA_word)

	for row in range(0, table_rows_len):
		for cell in table.rows[row].cells:
			cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
			cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

	doc = helper_funcs.set_autofit(doc)

	doc.add_paragraph("R squared adjusted = {R}".format(R="{:.2f}".format(output_df["R2adj"][0])))
	doc.add_paragraph("Dependent Variable: {DV}".format(DV=global_vars.raw_mr_outcomevar))
	helper_funcs.add_correction_message_word(doc)

	helper_funcs.savefile(doc=doc)