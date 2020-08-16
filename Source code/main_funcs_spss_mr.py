# MUltiple tests corrections and FOrmatted tables Software (MUFOS)
# Copyright (C) 2020  Nikolay Petrov, Vasil Atanasov, & Trevor Thompson
import global_vars
import helper_funcs

import pandas as pd
import numpy as np
from openpyxl import Workbook
from openpyxl.styles import Border, Side, Alignment, Font
from openpyxl.utils import get_column_letter
from openpyxl.utils.dataframe import dataframe_to_rows
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

#-----------------------------------------------------------Modified raw data dataframe----------------------------------------------------
def spss_mr_generate_mod_raw_data_df(raw_data_df):
	if not(raw_data_df.columns[0] == "Coefficientsa" or
		raw_data_df.columns[1:].str.contains("Unnamed:").all() or
		list(raw_data_df.iloc[0, 0:7]) == ["Model", np.nan, "Unstandardized Coefficients", np.nan, "Standardized Coefficients", "t", "Sig."] or
		list(raw_data_df.iloc[1, 0:7]) == [np.nan, np.nan, "B", "Std. Error", "Beta", np.nan, np.nan] or
		"Dependent Variable:" in raw_data_df.iloc[raw_data_df.index[-1], 0] or
		raw_data_df.iloc[2, 1] == "(Constant)"):
		
		raise Exception("The provided SPSS regression table does not seem to be in the accepted format. Please use the \"How to export SPSS table?\" button for guidance or refer to the documentation.")

	mod_raw_data_df = raw_data_df.copy()

	# renaming columns as the multi-level formatting of the spss table makes it difficult to work with
	renaming_dict = {}
	for ind, col_name in enumerate(mod_raw_data_df.columns):
		if "Unnamed" in col_name:
			if not pd.isna(mod_raw_data_df.iloc[0, ind]):
				renaming_dict[col_name] = mod_raw_data_df.iloc[0, ind]
			elif not pd.isna(mod_raw_data_df.iloc[1, ind]):
				renaming_dict[col_name] = mod_raw_data_df.iloc[1, ind]

	mod_raw_data_df = mod_raw_data_df.rename(columns=renaming_dict).drop([mod_raw_data_df.index[0], mod_raw_data_df.index[1]]).reset_index(drop=True)

	return mod_raw_data_df

#-----------------------------------------------------------Output dataframe----------------------------------------------------
def spss_mr_generate_output_df(mod_raw_data_df):
	variables_list = ["(Constant)"] + list(mod_raw_data_df[mod_raw_data_df.columns[1]])[1:-1]

	try:
		ci_low_list = ["{:.2f}".format(x) for x in list(mod_raw_data_df["95.0% Confidence Interval for B"])[:-1]]
	except KeyError:
		ci_low_list = [""] * (len(variables_list))
	try:
		ci_high_list = ["{:.2f}".format(x) for x in list(mod_raw_data_df["Upper Bound"])[:-1]]
	except KeyError:
		ci_high_list = [""] * (len(variables_list))
	beta_list = ["{:.2f}".format(x) for x in list(mod_raw_data_df["Standardized Coefficients"])[:-1]]

	output_df = pd.DataFrame()
	output_df["Variable"] = variables_list
	output_df["B"] = ["{:.2f}".format(x) for x in list(mod_raw_data_df["Unstandardized Coefficients"])[:-1]]
	output_df["95% CI"] = ["["+low+", "+high+"]" for low,high in zip(ci_low_list, ci_high_list)]
	output_df["beta"] = ['' if x=='nan' else x for x in beta_list] # the beta for the constant is missing from the spss output
	output_df["t"] = ["{:.2f}".format(x) for x in list(mod_raw_data_df["t"])[:-1]]
	output_df["pvalues"] = list(mod_raw_data_df["Sig."])[:-1] # there's a trailing np.nan at the end of the list

	return output_df

#-----------------------------------------------------------Saving data----------------------------------------------------
def spss_mr_apa_table_excel(mod_raw_data_df, output_df):
	output_df.drop(columns=["pvalues"], inplace=True)
	pd.options.mode.chained_assignment = None
	output_df["adjusted_pvalues"] = output_df["adjusted_pvalues"].map(helper_funcs.pvalue_formatting)
	output_df.rename(columns = {"adjusted_pvalues": "p"}, inplace=True)

	wb = Workbook()
	ws = wb.active

	for row in dataframe_to_rows(output_df, index=False, header=True):
		ws.append(row)

	for cell in ws[1]:
		cell.font = global_vars.font_header
		
	for row in range(1, len(output_df)+2):
		for cell in ws[row]:
			cell.alignment = global_vars.alignment_center
	
	for cell in ws[1]:
		cell.border = Border(bottom=global_vars.border_APA, top=global_vars.border_APA)

	for cell in ws[len(output_df)+1]:
		cell.border = Border(bottom=global_vars.border_APA)

	DV_cell = mod_raw_data_df[mod_raw_data_df.columns[0]][len(mod_raw_data_df[mod_raw_data_df.columns[0]])-1]
	DV = DV_cell[DV_cell.find(":")+2:]
	table_notes = ["R squared adjusted = X.XX", "Dependent variable: {}".format(DV)]
	if output_df["95% CI"][0] == "[,]":
		table_notes.append("Confidence intervals were not found in the SPSS table. Please add them to your SPSS table and re-run the program or add them manually.")

	helper_funcs.add_table_notes(ws, table_notes)

	helper_funcs.savefile(wb=wb)

def spss_mr_apa_table_word(mod_raw_data_df, output_df):
	output_df.drop(columns=["pvalues"], inplace=True)
	pd.options.mode.chained_assignment = None
	output_df["adjusted_pvalues"] = output_df["adjusted_pvalues"].map(helper_funcs.pvalue_formatting)
	output_df.rename(columns = {"adjusted_pvalues": "p"}, inplace=True)

	doc = Document()

	table_rows_len = len(output_df) + 1
	table_cols_len = len(output_df.columns)
	table = doc.add_table(rows=table_rows_len, cols=table_cols_len)

	for ind, var in enumerate(output_df.columns):
		table.cell(row_idx=0, col_idx=ind).text = var

	for row in range(1, table_rows_len):
		for col in range(0, table_cols_len):
			table.cell(row_idx=row, col_idx=col).text = output_df.iloc[row-1, col]

	for cell in table.rows[0].cells:
		helper_funcs.word_style(cell, italic=True)

	for row in range(0, table_rows_len):
		for cell in table.rows[row].cells:
			cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
			cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

	for cell in table.rows[0].cells:
		helper_funcs.set_cell_border(cell, top=global_vars.border_APA_word, bottom=global_vars.border_APA_word)
	for cell in table.rows[table_rows_len-1].cells:
		helper_funcs.set_cell_border(cell, bottom=global_vars.border_APA_word)

	doc = helper_funcs.set_autofit(doc)

	DV_cell = mod_raw_data_df[mod_raw_data_df.columns[0]][len(mod_raw_data_df[mod_raw_data_df.columns[0]])-1]
	DV = DV_cell[DV_cell.find(":")+2:]
	doc.add_paragraph("R squared adjusted = X.XX")
	doc.add_paragraph("Dependent Variable: {}".format(DV))
	if output_df["95% CI"][0] == "[,]":
		doc.add_paragraph("Confidence intervals were not found in the SPSS table. Please add them to your SPSS table and re-run the program or add them manually.")
	helper_funcs.add_correction_message_word(doc)

	helper_funcs.savefile(doc=doc)