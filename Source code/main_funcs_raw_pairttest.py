# MUltiple tests corrections and FOrmatted tables Software (MUFOS)
# Copyright (C) 2020  Nikolay Petrov, Vasil Atanasov, & Trevor Thompson
import global_vars
import helper_funcs

import pandas as pd
import numpy as np
from scipy import stats
import researchpy as rp
from openpyxl import Workbook
from openpyxl.styles import Border, Side, Alignment, Font
from openpyxl.utils import get_column_letter
from openpyxl.utils.dataframe import dataframe_to_rows
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

def raw_pairttest_generate_output_df(mod_raw_data_df):
	effect_size_df_row_lookup = {"Cohen's d": 6, "Hedge's g": 7, "Glass's delta": 8}
	dictionaries_list=[]

	for pair in global_vars.raw_pairttest_var_pairs:
		series_time1 = mod_raw_data_df[pair[0]][(mod_raw_data_df[pair[0]].notnull()) & (mod_raw_data_df[pair[1]].notnull())] #this also seems to be done withi nthe researchpy ttest func but it's fine
		series_time2 = mod_raw_data_df[pair[1]][(mod_raw_data_df[pair[0]].notnull()) & (mod_raw_data_df[pair[1]].notnull())]

		result = rp.ttest(series_time1, series_time2, group1_name=pair[0], group2_name=pair[1],
							equal_variances=True, 
							paired=True)
		
		ttest_stats_df = result[1]

		current_dict={}
		current_dict["Variable"] = "{var1} - {var2}".format(var1=pair[0], var2=pair[1])
		current_dict["Time1_N"] = series_time1.count()
		current_dict["Time1_Mean"] = np.mean(series_time1)
		current_dict["Time1_SD"] = np.std(series_time1)
		current_dict["Time2_N"] = series_time2.count()
		current_dict["Time2_Mean"] = np.mean(series_time2)
		current_dict["Time2_SD"] = np.std(series_time2)
		current_dict["Degrees of Freedom"] = ttest_stats_df.iloc[1, 1]
		current_dict["t"] = ttest_stats_df.iloc[2, 1]
		current_dict[global_vars.effect_size_choice] = np.nan if global_vars.effect_size_choice == "None" else ttest_stats_df.iloc[effect_size_df_row_lookup[global_vars.effect_size_choice], 1]
		current_dict["pvalues"] = ttest_stats_df.iloc[3, 1]
		
		dictionaries_list.append(current_dict)

	output_df = pd.DataFrame(dictionaries_list)

	return output_df
		
def raw_pairttest_apa_table_excel(mod_raw_data_df, output_df):
	output_df.drop(columns = ["pvalues"], inplace=True)
	apa_table_df = output_df[["Variable", "Time1_Mean", "Time1_SD", "Time2_Mean", "Time2_SD", "Degrees of Freedom", "t", global_vars.effect_size_choice, "adjusted_pvalues"]]

	pd.options.mode.chained_assignment = None
	apa_table_df[list(apa_table_df.columns)[1:-1]] = apa_table_df[list(apa_table_df.columns)[1:-1]].applymap(lambda x: "{:.2f}".format(x))

	apa_table_df["adjusted_pvalues"] = apa_table_df["adjusted_pvalues"].map(helper_funcs.pvalue_formatting)
	pd.options.mode.chained_assignment = "warn"

	wb = Workbook()
	ws = wb.active
	
	ws.cell(row=1, column=1).value = "Variable"
	ws.merge_cells("A1:A2")
	ws.cell(row=1, column=1).font = global_vars.font_header
	
	ws.cell(row=1, column=2).value = "Time 1"
	ws.merge_cells("B1:C1")
	
	ws.cell(row=1, column=4).value = "Time 2"
	ws.merge_cells("D1:E1")
	
	ws.cell(row=1, column=6).value = "df"
	ws.merge_cells("F1:F2")
	ws.cell(row=1, column=6).font = global_vars.font_header
	
	ws.cell(row=1, column=7).value = "t"
	ws.merge_cells("G1:G2")
	ws.cell(row=1, column=7).font = global_vars.font_header
	
	ws.cell(row=1, column=8).value = global_vars.effect_size_choice
	ws.merge_cells("H1:H2")
	ws.cell(row=1, column=8).font = global_vars.font_header
	
	ws.cell(row=1, column=9).value = "p"
	ws.merge_cells("I1:I2")
	ws.cell(row=1, column=9).font = global_vars.font_header
	
	for col in range(2,5,2):
		ws.cell(row=2, column=col).value = "M"
		ws.cell(row=2, column=col).font = global_vars.font_header
		ws.cell(row=2, column=col+1).value = "SD"
		ws.cell(row=2, column=col+1).font = global_vars.font_header
	
	for row in dataframe_to_rows(apa_table_df, index=False, header=False):
		ws.append(row)

	for row in range(1, len(apa_table_df)+3):
		for cell in ws[row]:
			cell.alignment = global_vars.alignment_center
	
	for cell in ws[2] + ws[len(apa_table_df)+2]:
		cell.border = Border(bottom=global_vars.border_APA)

	for cell in ws[1]:
		cell.border = Border(top=global_vars.border_APA)

	if global_vars.effect_size_choice == "None":
		ws.delete_cols(8)

	helper_funcs.add_table_notes(ws, [])

	helper_funcs.savefile(wb=wb)

def raw_pairttest_apa_table_word(mod_raw_data_df, output_df):
	output_df.drop(columns = ["pvalues"], inplace=True)
	apa_table_df = output_df[["Variable", "Time1_Mean", "Time1_SD", "Time2_Mean", "Time2_SD", "Degrees of Freedom", "t", global_vars.effect_size_choice, "adjusted_pvalues"]]

	pd.options.mode.chained_assignment = None
	apa_table_df[list(apa_table_df.columns)[1:-1]] = apa_table_df[list(apa_table_df.columns)[1:-1]].applymap(lambda x: "{:.2f}".format(x))

	apa_table_df["adjusted_pvalues"] = apa_table_df["adjusted_pvalues"].map(helper_funcs.pvalue_formatting)
	pd.options.mode.chained_assignment = "warn"

	doc = Document()
	table_rows_len = len(apa_table_df) + 2
	table_cols_len = len(apa_table_df.columns)
	table = doc.add_table(rows=table_rows_len, cols=table_cols_len)

	table.cell(row_idx=0, col_idx=0).text = "Variable"
	table.cell(row_idx=0, col_idx=0).merge(table.cell(row_idx=1, col_idx=0))
	helper_funcs.word_style(table.cell(row_idx=0, col_idx=0), italic=True)

	table.cell(row_idx=0, col_idx=1).text = "Time 1"
	table.cell(row_idx=0, col_idx=1).merge(table.cell(row_idx=0, col_idx=2))
	
	table.cell(row_idx=0, col_idx=3).text = "Time 2"
	table.cell(row_idx=0, col_idx=3).merge(table.cell(row_idx=0, col_idx=4))

	table.cell(row_idx=0, col_idx=5).text = "df"
	table.cell(row_idx=0, col_idx=5).merge(table.cell(row_idx=1, col_idx=5))
	helper_funcs.word_style(table.cell(row_idx=0, col_idx=5), italic=True)
	
	table.cell(row_idx=0, col_idx=6).text = "t"
	table.cell(row_idx=0, col_idx=6).merge(table.cell(row_idx=1, col_idx=6))
	helper_funcs.word_style(table.cell(row_idx=0, col_idx=6), italic=True)
	
	table.cell(row_idx=0, col_idx=7).text = global_vars.effect_size_choice
	if global_vars.effect_size_choice != "None": # otherwise cant remove with delete columns below if merged; see helper_funcs.delete_columns_word func notes
		table.cell(row_idx=0, col_idx=7).merge(table.cell(row_idx=1, col_idx=7))
	helper_funcs.word_style(table.cell(row_idx=0, col_idx=7), italic=True)
	
	table.cell(row_idx=0, col_idx=8).text = "p"
	table.cell(row_idx=0, col_idx=8).merge(table.cell(row_idx=1, col_idx=8))
	helper_funcs.word_style(table.cell(row_idx=0, col_idx=8), italic=True)

	for col in range(1, 4, 2):
		table.cell(row_idx=1, col_idx=col).text = "M"
		helper_funcs.word_style(table.cell(row_idx=1, col_idx=col), italic=True)
		table.cell(row_idx=1, col_idx=col+1).text = "SD"
		helper_funcs.word_style(table.cell(row_idx=1, col_idx=col+1), italic=True)
	
	for row in range(2, table_rows_len):
		for col in range(0, table_cols_len):
			table.cell(row_idx=row, col_idx=col).text = apa_table_df.iloc[row-2, col]

	for row in range(0, table_rows_len):
		for cell in table.rows[row].cells:
			cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
			cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

	for cell in table.rows[0].cells:
		helper_funcs.set_cell_border(cell, top=global_vars.border_APA_word)
	for cell in table.rows[2].cells:
		helper_funcs.set_cell_border(cell, top=global_vars.border_APA_word)
	for cell in table.rows[table_rows_len-1].cells:
		helper_funcs.set_cell_border(cell, bottom=global_vars.border_APA_word)

	if global_vars.effect_size_choice == "None":
		helper_funcs.delete_columns_word(table, [7])

	doc = helper_funcs.set_autofit(doc)

	helper_funcs.add_correction_message_word(doc)

	helper_funcs.savefile(doc=doc)