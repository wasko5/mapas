# MUltiple tests corrections and FOrmatted tables Software (MUFOS)
# Copyright (C) 2020  Nikolay Petrov, Vasil Atanasov, & Trevor Thompson
import global_vars
import helper_funcs

import pandas as pd
import numpy as np
from scipy import stats
from openpyxl import Workbook
from openpyxl.styles import Border, Side, Alignment, Font
from openpyxl.utils import get_column_letter
from openpyxl.utils.dataframe import dataframe_to_rows
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

#-----------------------------------------------------------Modified raw data dataframe----------------------------------------------------
def summ_inddtest_generate_mod_raw_data_df(raw_data_df, numeric_cols, non_numeric_cols=[]):
	mod_raw_data_df = raw_data_df[numeric_cols + non_numeric_cols] # does NOT raise errors when array is blank (default arg)
	mod_raw_data_df[numeric_cols] = mod_raw_data_df[numeric_cols].apply(pd.to_numeric, errors="coerce") # non-numeric values will be np.nan
	try:
		mod_raw_data_df[global_vars.summ_indttest_var] = mod_raw_data_df[global_vars.summ_indttest_var].astype(str) # does NOT raise errors when array is blank (default arg)
	except:
		raise Exception("The data could not be processed. Please ensure that the non-numeric columns contain only strings.")
	
	if global_vars.summ_indttest_equal_var == "":
		mod_raw_data_df["Equal Variances"] = [True] * len(mod_raw_data_df)
	else:
		if not raw_data_df[global_vars.summ_indttest_equal_var].isin([True, False]).all():
			raise Exception("A value different from TRUE or FALSE is found in column \'{c}\'.\nPlease ensure that only TRUE and FALSE values are provided.\nFor more information, see the software documentation at {d}, \"Input types\" section".format(c=global_vars.summ_indttest_equal_var, d=global_vars.software_page))
		
	return mod_raw_data_df

#-----------------------------------------------------------Output dataframe----------------------------------------------------
def summ_indttest_generate_output_df(mod_raw_data_df):
	if global_vars.summ_indttest_equal_var != "":
		equal_var_col = global_vars.summ_indttest_equal_var
	else:
		equal_var_col = list(mod_raw_data_df.columns)[-1]

	output_dict = {global_vars.summ_indttest_var: [], global_vars.summ_indttest_meanOne: [], global_vars.summ_indttest_sdOne: [], global_vars.summ_indttest_meanTwo: [],
					global_vars.summ_indttest_sdTwo: [], "Degrees of Freedom": [], "t": [], global_vars.effect_size_choice: [], "pvalues": []}

	for row in range(0, len(mod_raw_data_df)):
		current_series = mod_raw_data_df.loc[row]
		n1 = current_series[global_vars.summ_indttest_nOne]
		n2 = current_series[global_vars.summ_indttest_nTwo]

		output_dict[global_vars.summ_indttest_var].append(current_series[global_vars.summ_indttest_var])
		output_dict[global_vars.summ_indttest_meanOne].append(current_series[global_vars.summ_indttest_meanOne])
		output_dict[global_vars.summ_indttest_sdOne].append(current_series[global_vars.summ_indttest_sdOne])
		output_dict[global_vars.summ_indttest_meanTwo].append(current_series[global_vars.summ_indttest_meanTwo])
		output_dict[global_vars.summ_indttest_sdTwo].append(current_series[global_vars.summ_indttest_sdTwo])

		if current_series[equal_var_col]: # equal vars assumed
			deg_free = n1 + n2 - 2
		else:
			s1_pooled = current_series[global_vars.summ_indttest_sdOne]**2
			s2_pooled = current_series[global_vars.summ_indttest_sdTwo]**2
			deg_free = (s1_pooled/n1 + s2_pooled/n2)**2 / ( (1/(n1-1)*(s1_pooled/n1)**2) + (1/(n2-1)*(s2_pooled/n2)**2) )
		output_dict["Degrees of Freedom"].append(deg_free)
		# stats.ttest_ind_from_stats itself decides whether to to perform student's t-test or welch's based on equal_var
		t, p = stats.ttest_ind_from_stats(mean1 = current_series[global_vars.summ_indttest_meanOne], std1 = current_series[global_vars.summ_indttest_sdOne], 
										nobs1 = n1, mean2 = current_series[global_vars.summ_indttest_meanTwo], std2 = current_series[global_vars.summ_indttest_sdTwo], 
										nobs2 = n2, equal_var=current_series[equal_var_col])
		output_dict["t"].append(t)
		output_dict["pvalues"].append(p)

		effect_size = helper_funcs.calc_ttest_effect_size(effect_size_choice = global_vars.effect_size_choice, t=t, n1=n1, n2=n2,
											m1=current_series[global_vars.summ_indttest_meanOne], m2=current_series[global_vars.summ_indttest_meanTwo],
											sd1=current_series[global_vars.summ_indttest_sdOne])
		output_dict[global_vars.effect_size_choice].append(effect_size)

	output_df = pd.DataFrame(output_dict)
	
	return output_df

#-----------------------------------------------------------Saving data----------------------------------------------------
def summ_indttest_apa_table_excel(mod_raw_data_df, output_df):
	output_df.drop(columns = ["pvalues"], inplace=True)

	pd.options.mode.chained_assignment = None
	output_df[list(output_df.columns)[1:-1]] = output_df[list(output_df.columns)[1:-1]].applymap(lambda x: "{:.2f}".format(x))

	output_df["adjusted_pvalues"] = output_df["adjusted_pvalues"].map(helper_funcs.pvalue_formatting)
	pd.options.mode.chained_assignment = "warn"

	wb = Workbook()
	ws = wb.active

	ws.cell(row=1, column=1).value = "Variable"
	ws.merge_cells("A1:A2")
	ws.cell(row=1, column=1).font = global_vars.font_header
	
	ws.cell(row=1, column=2).value = "Group1, n={n}".format(n=mod_raw_data_df[global_vars.summ_indttest_nOne][0])
	ws.merge_cells("B1:C1")
	
	ws.cell(row=1, column=4).value = "Group2, n={n}".format(n=mod_raw_data_df[global_vars.summ_indttest_nTwo][0])
	ws.merge_cells("D1:E1")
	
	ws.cell(row=1, column=6).value = "df"
	ws.merge_cells("F1:F2")
	ws.cell(row=1, column=6).font = global_vars.font_header
	
	ws.cell(row=1, column=7).value = "t"
	ws.merge_cells("G1:G2")
	ws.cell(row=1, column=7).font = global_vars.font_header
	
	ws.cell(row=1, column=8).value = global_vars.effect_size_choice
	ws.merge_cells("H1:H2")
	ws.cell(row=1, column=8).font = global_vars.font_header
	
	ws.cell(row=1, column=9).value = "p"
	ws.merge_cells("I1:I2")
	ws.cell(row=1, column=9).font = global_vars.font_header
	
	for col in range(2,5,2):
		ws.cell(row=2, column=col).value = "M"
		ws.cell(row=2, column=col).font = global_vars.font_header
		ws.cell(row=2, column=col+1).value = "SD"
		ws.cell(row=2, column=col+1).font = global_vars.font_header
	
	for row in dataframe_to_rows(output_df, index=False, header=False):
		ws.append(row)

	for row in range(1, len(output_df)+3):
		for cell in ws[row]:
			cell.alignment = global_vars.alignment_center

	for cell in ws[1]:
		cell.border = Border(top=global_vars.border_APA)
	for cell in ws[2] + ws[len(output_df)+2]:
		cell.border = Border(bottom=global_vars.border_APA)

	if global_vars.effect_size_choice == "None":
		ws.delete_cols(8)

	helper_funcs.add_table_notes(ws, [])

	helper_funcs.savefile(wb=wb)

def summ_indttest_apa_table_word(mod_raw_data_df, output_df):
	output_df.drop(columns = ["pvalues"], inplace=True)

	pd.options.mode.chained_assignment = None
	output_df[list(output_df.columns)[1:-1]] = output_df[list(output_df.columns)[1:-1]].applymap(lambda x: "{:.2f}".format(x))

	output_df["adjusted_pvalues"] = output_df["adjusted_pvalues"].map(helper_funcs.pvalue_formatting)
	pd.options.mode.chained_assignment = "warn"

	doc = Document()
	table_rows_len = len(output_df) + 2
	table_cols_len = len(output_df.columns)
	table = doc.add_table(rows=table_rows_len, cols=table_cols_len)

	table.cell(row_idx=0, col_idx=0).text = "Variable"
	table.cell(row_idx=0, col_idx=0).merge(table.cell(row_idx=1, col_idx=0))
	helper_funcs.word_style(table.cell(row_idx=0, col_idx=0), italic=True)

	table.cell(row_idx=0, col_idx=1).text = "Group1, n={n}".format(n=mod_raw_data_df[global_vars.summ_indttest_nOne][0])
	table.cell(row_idx=0, col_idx=1).merge(table.cell(row_idx=0, col_idx=2))

	table.cell(row_idx=0, col_idx=3).text = "Group2, n={n}".format(n=mod_raw_data_df[global_vars.summ_indttest_nTwo][0])
	table.cell(row_idx=0, col_idx=3).merge(table.cell(row_idx=0, col_idx=4))

	table.cell(row_idx=0, col_idx=5).text = "df"
	table.cell(row_idx=0, col_idx=5).merge(table.cell(row_idx=1, col_idx=5))
	helper_funcs.word_style(table.cell(row_idx=0, col_idx=5), italic=True)

	table.cell(row_idx=0, col_idx=6).text = "t"
	table.cell(row_idx=0, col_idx=6).merge(table.cell(row_idx=1, col_idx=6))
	helper_funcs.word_style(table.cell(row_idx=0, col_idx=6), italic=True)

	table.cell(row_idx=0, col_idx=7).text = global_vars.effect_size_choice
	if global_vars.effect_size_choice != "None": # otherwise cant remove with delete columns below if merged; see helper_funcs.delete_columns_word func notes
		table.cell(row_idx=0, col_idx=7).merge(table.cell(row_idx=1, col_idx=7))
	helper_funcs.word_style(table.cell(row_idx=0, col_idx=7), italic=True)
	
	table.cell(row_idx=0, col_idx=8).text = "p"
	table.cell(row_idx=0, col_idx=8).merge(table.cell(row_idx=1, col_idx=8))
	helper_funcs.word_style(table.cell(row_idx=0, col_idx=8), italic=True)

	for col in range(1, 4, 2):
		table.cell(row_idx=1, col_idx=col).text = "M"
		helper_funcs.word_style(table.cell(row_idx=1, col_idx=col), italic=True)
		table.cell(row_idx=1, col_idx=col+1).text = "SD"
		helper_funcs.word_style(table.cell(row_idx=1, col_idx=col+1), italic=True)

	for row in range(2, table_rows_len):
		for col in range(0, table_cols_len):
			table.cell(row_idx=row, col_idx=col).text = output_df.iloc[row-2, col]

	for row in range(0, table_rows_len):
		for cell in table.rows[row].cells:
			cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
			cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

	for cell in table.rows[0].cells:
		helper_funcs.set_cell_border(cell, top=global_vars.border_APA_word)
	for cell in table.rows[2].cells:
		helper_funcs.set_cell_border(cell, top=global_vars.border_APA_word)
	for cell in table.rows[table_rows_len-1].cells:
		helper_funcs.set_cell_border(cell, bottom=global_vars.border_APA_word)
	
	if global_vars.effect_size_choice == "None":
		helper_funcs.delete_columns_word(table, [7])

	doc = helper_funcs.set_autofit(doc)

	helper_funcs.add_correction_message_word(doc)

	helper_funcs.savefile(doc=doc)